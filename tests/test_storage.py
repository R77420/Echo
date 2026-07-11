"""Tests du module storage.py (persistance JSON + génération .docx).

Tous les tests utilisent tmp_path : aucun vrai fichier n'est touché.
Lancer : pytest tests/test_storage.py -v
"""

import datetime
import os
import sys

import pytest

# Permet d'importer storage.py situé à la racine du projet.
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import storage  # noqa: E402


# --------------------------------------------------------------------------- #
#  Fixtures / helpers
# --------------------------------------------------------------------------- #

def _chemin(tmp_path):
    return os.path.join(tmp_path, "consultations.json")


def _record(cid="a1", nom="DUPONT", file_path="", summary=""):
    return {
        "id": cid,
        "date": "2026-06-13T09:00:00",
        "patient": {"nom": nom, "prenom": "Marie",
                    "naissance": "12/03/1980", "motif": "Toux"},
        "summary": summary,
        "file_path": file_path,
        "duration_min": 12,
    }


# --------------------------------------------------------------------------- #
#  JSON : écriture / lecture
# --------------------------------------------------------------------------- #

def test_ecrire_lire_consultation(tmp_path):
    chemin = _chemin(tmp_path)
    storage.ajouter_consultation(chemin, _record("a1", "DUPONT"))
    storage.ajouter_consultation(chemin, _record("b2", "MARTIN"))

    data = storage.charger_consultations(chemin)
    assert len(data) == 2
    # ajouter_consultation insère en tête → b2 d'abord.
    assert data[0]["id"] == "b2"
    assert data[1]["id"] == "a1"
    assert data[1]["patient"]["nom"] == "DUPONT"
    assert data[0]["patient"]["motif"] == "Toux"


def test_charger_consultations_absent(tmp_path):
    # Fichier inexistant → liste vide, pas d'erreur.
    assert storage.charger_consultations(_chemin(tmp_path)) == []


# --------------------------------------------------------------------------- #
#  Suppression
# --------------------------------------------------------------------------- #

def test_supprimer_consultation_historique_seul(tmp_path):
    chemin = _chemin(tmp_path)
    # Crée un faux .docx sur le disque.
    docx = os.path.join(tmp_path, "cr.docx")
    with open(docx, "w", encoding="utf-8") as f:
        f.write("contenu word")

    storage.ajouter_consultation(chemin, _record("a1", file_path=docx))
    storage.ajouter_consultation(chemin, _record("b2"))

    rec = storage.supprimer_consultation(chemin, "a1")

    assert rec is not None and rec["id"] == "a1"
    # L'entrée json a disparu...
    ids = [c["id"] for c in storage.charger_consultations(chemin)]
    assert ids == ["b2"]
    # ...mais le .docx reste sur le disque.
    assert os.path.isfile(docx)


def test_supprimer_consultation_definitif(tmp_path):
    chemin = _chemin(tmp_path)
    docx = os.path.join(tmp_path, "cr.docx")
    with open(docx, "w", encoding="utf-8") as f:
        f.write("contenu word")

    storage.ajouter_consultation(chemin, _record("a1", file_path=docx))

    rec = storage.supprimer_consultation_avec_fichier(chemin, "a1")

    assert rec is not None and rec["id"] == "a1"
    # .docx ET entrée json ont disparu.
    assert storage.charger_consultations(chemin) == []
    assert not os.path.exists(docx)


def test_suppression_definitive_fichier_absent(tmp_path):
    # Le .docx a déjà été déplacé/supprimé manuellement → pas d'erreur,
    # l'entrée est quand même retirée.
    chemin = _chemin(tmp_path)
    storage.ajouter_consultation(
        chemin, _record("a1", file_path=os.path.join(tmp_path, "introuvable.docx")))

    rec = storage.supprimer_consultation_avec_fichier(chemin, "a1")
    assert rec is not None
    assert storage.charger_consultations(chemin) == []


def test_suppression_id_inexistant(tmp_path):
    chemin = _chemin(tmp_path)
    storage.ajouter_consultation(chemin, _record("a1"))

    # Ne doit pas lever, renvoie None, et ne touche pas aux entrées existantes.
    rec = storage.supprimer_consultation(chemin, "zzz")
    assert rec is None
    assert len(storage.charger_consultations(chemin)) == 1


# --------------------------------------------------------------------------- #
#  Mise à jour du résumé
# --------------------------------------------------------------------------- #

def test_maj_resume(tmp_path):
    chemin = _chemin(tmp_path)
    storage.ajouter_consultation(chemin, _record("a1", summary=""))
    storage.ajouter_consultation(chemin, _record("b2", summary="ancien"))

    storage.maj_consultation_resume(chemin, "a1", "Nouveau résumé patient")

    data = {c["id"]: c for c in storage.charger_consultations(chemin)}
    # Le résumé ciblé est mis à jour...
    assert data["a1"]["summary"] == "Nouveau résumé patient"
    # ...les autres champs de a1 restent intacts...
    assert data["a1"]["patient"]["nom"] == "DUPONT"
    assert data["a1"]["duration_min"] == 12
    # ...et les autres entrées ne bougent pas.
    assert data["b2"]["summary"] == "ancien"


def test_maj_resume_id_inexistant(tmp_path):
    chemin = _chemin(tmp_path)
    storage.ajouter_consultation(chemin, _record("a1", summary="x"))
    # id inconnu → pas d'erreur, rien ne change.
    storage.maj_consultation_resume(chemin, "zzz", "ignoré")
    assert storage.charger_consultations(chemin)[0]["summary"] == "x"


# --------------------------------------------------------------------------- #
#  Génération .docx
# --------------------------------------------------------------------------- #

def test_generation_docx_champs_optionnels(tmp_path):
    pytest.importorskip("docx")  # python-docx requis
    from docx import Document

    chemin = os.path.join(tmp_path, "cr.docx")
    now = datetime.datetime(2026, 6, 13, 9, 0)
    infos = {"nom": "DUPONT", "prenom": "", "naissance": "", "motif": ""}
    entries = [("09:00:00", "Medecin", "Bonjour"),
               ("09:00:05", "Patient", "Bonjour docteur")]

    # Ne doit pas crasher avec seulement le Nom.
    storage.ecrire_docx(chemin, infos, now, None, entries, annexes=[])
    assert os.path.isfile(chemin)

    doc = Document(chemin)
    texte_complet = "\n".join(p.text for p in doc.paragraphs)
    cellules = [c.text for t in doc.tables for r in t.rows for c in r.cells]

    # Le Nom est présent ; prénom/né(e) le omis (vides) ; Motif = "—".
    assert "DUPONT" in cellules
    assert "Prénom" not in cellules
    assert "Né(e) le" not in cellules
    assert "—" in cellules                      # motif vide → tiret
    assert "Compte-rendu de consultation" in texte_complet
    # La transcription intégrale figure bien.
    assert "Transcription intégrale" in texte_complet


def test_generation_docx_complet(tmp_path):
    pytest.importorskip("docx")
    from docx import Document

    chemin = os.path.join(tmp_path, "cr.docx")
    now = datetime.datetime(2026, 6, 13, 9, 0)
    infos = {"nom": "MARTIN", "prenom": "Paul",
             "naissance": "05/09/1975", "motif": "Suivi"}
    entries = [("10:00:00", "Patient", "J'ai mal à la tête")]

    storage.ecrire_docx(chemin, infos, now, None, entries)
    cellules = [c.text for t in Document(chemin).tables for r in t.rows for c in r.cells]
    assert "MARTIN" in cellules
    assert "Paul" in cellules
    assert "05/09/1975" in cellules
    assert "Suivi" in cellules


# --------------------------------------------------------------------------- #
#  Retry sur verrou Windows
# --------------------------------------------------------------------------- #

def test_retry_sur_lock(tmp_path, monkeypatch):
    """Simule un PermissionError aux 2 premières ouvertures en écriture,
    puis un succès → le retry doit aboutir."""
    chemin = _chemin(tmp_path)
    storage.ajouter_consultation(chemin, _record("a1"))
    storage.ajouter_consultation(chemin, _record("b2"))

    vrai_open = open
    etat = {"echecs_restants": 2}

    def open_capricieux(file, mode="r", *args, **kwargs):
        # Ne fait échouer que les ouvertures en écriture du fichier ciblé.
        if str(file) == chemin and "w" in mode and etat["echecs_restants"] > 0:
            etat["echecs_restants"] -= 1
            raise PermissionError("verrou simulé (OneDrive)")
        return vrai_open(file, mode, *args, **kwargs)

    # sleep neutralisé pour un test rapide.
    monkeypatch.setattr(storage.time, "sleep", lambda s: None)
    monkeypatch.setattr("builtins.open", open_capricieux)

    rec = storage.supprimer_consultation(chemin, "a1")

    # Les deux premiers essais ont échoué, le 3e a réussi.
    assert etat["echecs_restants"] == 0
    assert rec is not None and rec["id"] == "a1"

    # Restaure open pour la vérification finale.
    monkeypatch.undo()
    assert [c["id"] for c in storage.charger_consultations(chemin)] == ["b2"]


def test_retry_echoue_apres_3_tentatives(tmp_path, monkeypatch):
    """Si le verrou persiste, l'exception est propagée après 3 tentatives."""
    chemin = _chemin(tmp_path)
    storage.ajouter_consultation(chemin, _record("a1"))

    vrai_open = open

    def open_toujours_bloque(file, mode="r", *args, **kwargs):
        if str(file) == chemin and "w" in mode:
            raise PermissionError("verrou permanent")
        return vrai_open(file, mode, *args, **kwargs)

    monkeypatch.setattr(storage.time, "sleep", lambda s: None)
    monkeypatch.setattr("builtins.open", open_toujours_bloque)

    with pytest.raises(PermissionError):
        storage.supprimer_consultation(chemin, "a1")


# --------------------------------------------------------------------------- #
#  Patients : regroupement et recherche
# --------------------------------------------------------------------------- #

def _consult(cid, nom, prenom="", ddn="", date="2026-06-13T09:00:00"):
    return {
        "id": cid, "date": date,
        "patient": {"nom": nom, "prenom": prenom, "naissance": ddn,
                    "motif": ""},
        "summary": "", "file_path": "", "duration_min": 10,
    }


def test_extraire_patients_regroupe():
    consultations = [
        _consult("c1", "MARTIN", "Pierre", date="2026-06-01T09:00:00"),
        _consult("c2", "MARTIN", "Pierre", date="2026-06-10T09:00:00"),
        _consult("c3", "MARTIN", "Pierre", date="2026-06-20T09:00:00"),
    ]
    patients = storage.extraire_patients(consultations)
    assert len(patients) == 1
    p = patients[0]
    assert p["nb_consultations"] == 3
    assert sorted(p["consultation_ids"]) == ["c1", "c2", "c3"]
    assert p["derniere_consultation"] == "2026-06-20T09:00:00"


def test_patients_homonymes_distincts():
    consultations = [
        _consult("c1", "MARTIN", "Pierre"),
        _consult("c2", "MARTIN", "Sophie"),
    ]
    patients = storage.extraire_patients(consultations)
    assert len(patients) == 2


def test_meme_nom_sans_prenom():
    # Deux consultations "martin" sans prénom ni ddn → même clé, regroupées.
    consultations = [
        _consult("c1", "martin"),
        _consult("c2", "Martin"),
    ]
    patients = storage.extraire_patients(consultations)
    assert len(patients) == 1
    assert patients[0]["nb_consultations"] == 2


def test_regroupement_nom_seul():
    consultations = [
        _consult("c1", "martin", date="2026-06-01T09:00:00"),
        _consult("c2", "MARTIN", date="2026-06-02T09:00:00"),
        _consult("c3", "Martin", date="2026-06-03T09:00:00"),
    ]
    patients = storage.extraire_patients(consultations)
    assert len(patients) == 1
    assert patients[0]["nb_consultations"] == 3


def test_nom_seul_et_nom_prenom_distincts():
    # "MARTIN" sans prénom et "MARTIN Sophie" → 2 entrées distinctes.
    consultations = [
        _consult("c1", "MARTIN"),
        _consult("c2", "MARTIN", "Sophie"),
    ]
    patients = storage.extraire_patients(consultations)
    assert len(patients) == 2


def test_tri_derniere_consultation():
    consultations = [
        _consult("c1", "ANCIEN", "A", date="2026-01-01T09:00:00"),
        _consult("c2", "RECENT", "B", date="2026-06-01T09:00:00"),
    ]
    patients = storage.extraire_patients(consultations)
    assert patients[0]["nom"] == "RECENT"


def test_ddn_conservee_pour_affichage():
    # La ddn n'entre pas dans la clé mais est conservée si connue.
    consultations = [
        _consult("c1", "MARTIN", "Pierre"),
        _consult("c2", "MARTIN", "Pierre", ddn="12/03/1980"),
    ]
    patients = storage.extraire_patients(consultations)
    assert len(patients) == 1
    assert patients[0]["ddn"] == "12/03/1980"


def test_recherche_floue():
    consultations = [
        _consult("c1", "Martin", "Pierre"),
        _consult("c2", "Hélène", "Dupont"),
        _consult("c3", "Durand", "Hélène"),
    ]
    # Début de nom, insensible à la casse.
    res = storage.rechercher_patients(consultations, "mar")
    assert len(res) == 1 and res[0]["nom"] == "Martin"
    # Insensible aux accents : "héléne" trouve "Hélène" (nom ou prénom).
    res = storage.rechercher_patients(consultations, "héléne")
    noms = {(p["nom"], p["prenom"]) for p in res}
    assert ("Hélène", "Dupont") in noms
    assert ("Durand", "Hélène") in noms
    # Query vide → rien.
    assert storage.rechercher_patients(consultations, "") == []


def test_recherche_max_5():
    consultations = [
        _consult(f"c{i}", f"Martin{i}", "X") for i in range(8)
    ]
    assert len(storage.rechercher_patients(consultations, "mar")) == 5


def test_patient_manuel_sans_consultation(tmp_path):
    chemin = os.path.join(tmp_path, "patients.json")
    assert storage.ajouter_patient_manuel(chemin, "NOUVEAU", "Paul", "01/01/1990")
    # Doublon (même clé, casse/accents différents) refusé.
    assert not storage.ajouter_patient_manuel(chemin, "nouveau", "PAUL")
    # Nom vide refusé.
    assert not storage.ajouter_patient_manuel(chemin, "  ")

    manuels = storage.charger_patients_manuels(chemin)
    patients = storage.extraire_patients([], manuels)
    assert len(patients) == 1
    p = patients[0]
    assert p["nom"] == "NOUVEAU" and p["prenom"] == "Paul"
    assert p["ddn"] == "01/01/1990"
    assert p["nb_consultations"] == 0
    assert p["consultation_ids"] == []


def test_fusion_patient_manuel_et_consultations(tmp_path):
    chemin = os.path.join(tmp_path, "patients.json")
    storage.ajouter_patient_manuel(chemin, "Dupont", "Marie", "12/03/1980")
    manuels = storage.charger_patients_manuels(chemin)
    consultations = [
        _consult("c1", "DUPONT", "Marie", date="2026-07-01T09:00:00"),
        _consult("c2", "MARTIN", "Jean",  date="2026-07-02T09:00:00"),
    ]
    patients = storage.extraire_patients(consultations, manuels)
    # Une seule entrée pour Dupont Marie (fusion manuel + consultation).
    assert len(patients) == 2
    marie = next(p for p in patients if _normalise_test(p["nom"]) == "dupont")
    assert marie["nb_consultations"] == 1
    assert marie["consultation_ids"] == ["c1"]
    # La ddn saisie à la création manuelle est conservée.
    assert marie["ddn"] == "12/03/1980"
    # Patients avec consultation d'abord (tri par date), manuel seul en fin.
    storage.ajouter_patient_manuel(chemin, "SANSRDV", "Zoe")
    patients = storage.extraire_patients(
        consultations, storage.charger_patients_manuels(chemin))
    assert patients[-1]["nom"] == "SANSRDV"


def _normalise_test(s):
    return storage._normaliser(s)
