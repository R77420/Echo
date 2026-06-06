"""
Transcription temps reel d'une consultation Doctolib (Windows, 100 % local).

Capte simultanement :
  - le microphone du medecin
  - le son des haut-parleurs (= la voix du patient) via WASAPI loopback

Decoupe la parole aux silences (webrtcvad), transcrit chaque segment avec
faster-whisper, et affiche le texte dans une fenetre overlay toujours au premier
plan, a poser a cote de la fenetre Doctolib.

Rien ne quitte la machine : aucun appel reseau, aucun enregistrement sur disque.

------------------------------------------------------------------------------
DEPENDANCES
    pip install faster-whisper soundcard numpy webrtcvad-wheels
    (webrtcvad-wheels = meme API `import webrtcvad`, mais pre-compile pour
     Windows, donc pas besoin de Visual C++ Build Tools.)

SELECTION DU PERIPHERIQUE
    Par defaut on prend le loopback de la sortie Windows par defaut. Sur une
    machine avec routage virtuel (ex. SteelSeries Sonar), le defaut peut etre
    muet pour les medias. Au lancement, une LISTE DEROULANTE des sorties permet
    de choisir manuellement laquelle transcrire ; elle pre-selectionne
    PREFERRED_OUTPUT s'il correspond a une sortie disponible.
      - PREFERRED_OUTPUT : nom (fragment) pre-coche dans la liste. Sur cette
        machine = "Logitech PRO X" (sortie physique reelle derriere Sonar).
      - LOOPBACK_NAME : si renseigne, court-circuite la liste (mode sans GUI).
------------------------------------------------------------------------------
"""

import datetime
import json
import multiprocessing
import os
import queue
import re
import subprocess
import sys
import tempfile
import threading
import time
import tkinter as tk
import unicodedata
import uuid
from tkinter import filedialog, messagebox, ttk

import numpy as np
import soundcard as sc
import webrtcvad
from faster_whisper import WhisperModel

# ----------------------------- PARAMETRES ------------------------------------

MODEL_SIZE   = "large-v3-turbo"  # bench: RTF 0.95 sur CPU, meilleur que small sur vocab médical
DEVICE       = "cpu"       # "cuda" si carte NVIDIA disponible
COMPUTE_TYPE = "int8"      # "int8" sur CPU, "float16" sur GPU
LANGUAGE     = "fr"

# Noms des modèles (plus embarqués dans l'exe — téléchargés dans %APPDATA%\Echo\models\)
MODELE_WHISPER_DIR  = "faster-whisper-large-v3-turbo"
MODELE_EMBARQUE     = MODELE_WHISPER_DIR   # compatibilité ancienne constante

# Selection du loopback a transcrire (voix patient).
#   PREFERRED_OUTPUT : pre-coche dans la liste deroulante (None = sortie defaut).
#   LOOPBACK_NAME    : force sans afficher la liste (None = afficher la liste).
PREFERRED_OUTPUT = "Logitech PRO X"   # machine de dev avec Sonar -> sortie physique
LOOPBACK_NAME    = None
MIC_NAME         = None    # force le micro sans GUI (None = afficher la liste)

SAMPLE_RATE   = 16000      # impose par Whisper et webrtcvad
FRAME_MS      = 30         # webrtcvad accepte 10/20/30 ms
FRAME_SAMPLES = SAMPLE_RATE * FRAME_MS // 1000   # 480 echantillons
VAD_LEVEL     = 2          # 0 (permissif) .. 3 (agressif sur le bruit)
SILENCE_MS    = 700        # silence qui marque la fin d'un tour de parole
MIN_SPEECH_MS = 300        # ignore les bruits trop courts
MAX_SEG_MS    = 15000      # flush force pour ne pas accumuler indefiniment

stop_event = threading.Event()
segment_queue = queue.Queue()
display_queue = queue.Queue()

# ---- Gains de volume (thread-safe via GIL + lock en écriture) ----
_GAIN_PATIENT = 1.0
_GAIN_MIC     = 1.0
_gain_lock    = threading.Lock()

def _get_gain(label):
    return _GAIN_PATIENT if label == "Patient" else _GAIN_MIC


# ---- État VAD temps réel (qui parle maintenant) ----
# Lecture/écriture de booléens simples : atomique sous GIL, pas de lock requis.
_speaking_now = {"medecin": False, "patient": False}

def _set_speaking(label, is_speech):
    _speaking_now["patient" if label == "Patient" else "medecin"] = bool(is_speech)


# ----------------------------- CONFIG PERSISTANTE ----------------------------
# Principe : ne JAMAIS planter. Toute erreur de config est avalee -> defauts.

def dossier_config():
    """Dossier de configuration : %APPDATA%\\Echo (cree au besoin)."""
    base = os.environ.get("APPDATA") or os.path.expanduser("~")
    return os.path.join(base, "Echo")


def chemin_config():
    return os.path.join(dossier_config(), "config.json")


def charger_config():
    """Lit la config. Absente / illisible / corrompue -> {} (jamais d'erreur)."""
    try:
        with open(chemin_config(), "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def sauver_config(cfg):
    """Ecrit la config en best-effort. N'echoue jamais bruyamment."""
    try:
        os.makedirs(dossier_config(), exist_ok=True)
        with open(chemin_config(), "w", encoding="utf-8") as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def chemin_consultations():
    return os.path.join(dossier_config(), "consultations.json")

def charger_consultations():
    """Historique des consultations. Renvoie [] si absent/corrompu."""
    try:
        with open(chemin_consultations(), encoding="utf-8") as f:
            d = json.load(f)
        return d if isinstance(d, list) else []
    except Exception:
        return []

def ajouter_consultation(record):
    """Ajoute un enregistrement en tête de l'historique (best-effort)."""
    try:
        consultations = charger_consultations()
        consultations.insert(0, record)
        os.makedirs(dossier_config(), exist_ok=True)
        with open(chemin_consultations(), "w", encoding="utf-8") as f:
            json.dump(consultations, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def supprimer_consultation(cid):
    """Retire l'entrée d'id `cid` de l'historique et réécrit le fichier.

    Lecture et écriture se font toujours via `with open(...)` : aucun handle
    n'est conservé entre les appels. En cas de verrou Windows transitoire
    (OneDrive, antivirus, indexeur), on réessaie jusqu'à 3 fois.

    Renvoie l'enregistrement supprimé (dict) si trouvé, sinon None.
    Lève PermissionError/OSError si le fichier reste inaccessible après les
    tentatives.
    """
    chemin = chemin_consultations()
    os.makedirs(dossier_config(), exist_ok=True)
    derniere_exc = None
    for attempt in range(3):
        try:
            try:
                with open(chemin, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if not isinstance(data, list):
                    data = []
            except FileNotFoundError:
                data = []

            supprime = None
            restantes = []
            for c in data:
                if supprime is None and isinstance(c, dict) and c.get("id") == cid:
                    supprime = c
                else:
                    restantes.append(c)

            with open(chemin, "w", encoding="utf-8") as f:
                json.dump(restantes, f, ensure_ascii=False, indent=2)
            return supprime
        except (PermissionError, OSError) as exc:
            derniere_exc = exc
            if attempt < 2:
                time.sleep(0.15)
    raise derniere_exc

def message_simple(titre, message, genre="info"):
    """Affiche un message court (FR simple) sans dependre d'une fenetre existante."""
    try:
        r = tk.Tk()
        r.withdraw()
        r.attributes("-topmost", True)
        (messagebox.showerror if genre == "error" else messagebox.showinfo)(
            titre, message, parent=r)
        r.destroy()
    except Exception:
        pass


# ----------------------------- SELECTION PERIPHERIQUE ------------------------
# Resolveurs robustes : renvoient None plutot que de lever une exception.

def lister_sorties():
    try:
        return sc.all_speakers()
    except Exception:
        return []


def lister_micros():
    try:
        return sc.all_microphones()
    except Exception:
        return []


def nom_sortie_defaut():
    try:
        return str(sc.default_speaker().name)
    except Exception:
        return ""


def nom_micro_defaut():
    try:
        return str(sc.default_microphone().name)
    except Exception:
        return ""


def resoudre_loopback(nom):
    """Loopback dont le nom contient `nom`, sinon None (jamais d'exception)."""
    if not nom:
        return None
    try:
        for m in sc.all_microphones(include_loopback=True):
            if nom.lower() in m.name.lower():
                return m
    except Exception:
        pass
    return None


def loopback_defaut():
    """Loopback de la sortie Windows par defaut, sinon None."""
    return resoudre_loopback(nom_sortie_defaut())


def resoudre_micro(nom):
    """Entree reelle dont le nom contient `nom`, sinon None."""
    if not nom:
        return None
    try:
        for m in sc.all_microphones():
            if m.name == nom or nom.lower() in m.name.lower():
                return m
    except Exception:
        pass
    return None

def loopback_par_nom(fragment):
    """Renvoie le microphone-loopback dont le nom contient `fragment`."""
    micros = sc.all_microphones(include_loopback=True)
    for m in micros:
        if fragment.lower() in m.name.lower():
            return m
    raise RuntimeError("Loopback introuvable pour '%s'. Dispo : %s"
                       % (fragment, [m.name for m in micros]))


def micro_par_nom(fragment):
    """Renvoie l'entree reelle (sans loopback) dont le nom contient `fragment`."""
    micros = sc.all_microphones()  # entrees reelles uniquement
    for m in micros:
        if m.name == fragment or fragment.lower() in m.name.lower():
            return m
    raise RuntimeError("Micro introuvable pour '%s'. Dispo : %s"
                       % (fragment, [m.name for m in micros]))


def choisir_peripheriques():
    """Fenetre de reglages -> (loopback, micro_ou_None, dossier_sauvegarde) ou None.

    - Reglages persistants : pre-selection depuis config.json, ecriture au clic
      "Demarrer".
    - Robustesse (jamais de crash) :
        * aucune sortie audio  -> message clair + None (abandon propre) ;
        * aucun micro          -> on continue sans micro (patient seul) ;
        * peripherique memorise absent -> repli silencieux sur le defaut.
    - Mode dev : LOOPBACK_NAME / MIC_NAME court-circuitent la liste.
    """
    cfg = charger_config()
    dossier_sauv = cfg.get("dossier_sauvegarde")

    # Mode dev : constantes forcees (court-circuite la GUI si la sortie resout).
    if LOOPBACK_NAME or MIC_NAME:
        lb = resoudre_loopback(LOOPBACK_NAME) if LOOPBACK_NAME else loopback_defaut()
        mc = resoudre_micro(MIC_NAME) if MIC_NAME else None
        if lb is not None:
            return lb, mc, dossier_sauv

    sorties = lister_sorties()
    if not sorties:
        message_simple(
            "Aucune sortie audio",
            "Aucune sortie audio n'a été détectée sur cet ordinateur.\n\n"
            "Branchez un casque ou des haut-parleurs, puis relancez Écho.",
            genre="error")
        return None

    noms_sorties = [str(s.name) for s in sorties]
    noms_micros = [str(m.name) for m in lister_micros()]

    # Pre-selection sortie : config -> PREFERRED_OUTPUT -> defaut systeme.
    presel_sortie = None
    if cfg.get("sortie") in noms_sorties:
        presel_sortie = cfg["sortie"]
    elif PREFERRED_OUTPUT:
        for n in noms_sorties:
            if PREFERRED_OUTPUT.lower() in n.lower():
                presel_sortie = n
                break
    if not presel_sortie:
        d = nom_sortie_defaut()
        presel_sortie = d if d in noms_sorties else noms_sorties[0]

    # Pre-selection micro : config -> defaut systeme -> premier dispo.
    presel_micro = ""
    if noms_micros:
        if cfg.get("micro") in noms_micros:
            presel_micro = cfg["micro"]
        else:
            d = nom_micro_defaut()
            presel_micro = d if d in noms_micros else noms_micros[0]

    choix = {"sortie": presel_sortie, "micro": presel_micro, "valide": False}

    dlg = tk.Tk()
    dlg.title("Écho — réglages")
    dlg.configure(bg="#0d1117")
    dlg.attributes("-topmost", True)
    dlg.geometry("660x320+60+60")

    # --- Sortie (patient), libelle non-technique ---
    tk.Label(dlg, text="Ce que vous entendez (patient)",
             bg="#0d1117", fg="#79c0ff", font=("Segoe UI", 12, "bold")
             ).pack(anchor="w", padx=22, pady=(24, 4))
    var_sortie = tk.StringVar(value=presel_sortie)
    ttk.Combobox(dlg, textvariable=var_sortie, values=noms_sorties,
                 state="readonly", width=72).pack(anchor="w", padx=22)

    # --- Micro (medecin), libelle non-technique ---
    tk.Label(dlg, text="Votre micro (médecin)",
             bg="#0d1117", fg="#7ee787", font=("Segoe UI", 12, "bold")
             ).pack(anchor="w", padx=22, pady=(18, 4))
    var_micro = tk.StringVar(value=presel_micro if noms_micros else "(aucun micro détecté)")
    combo_micro = ttk.Combobox(
        dlg, textvariable=var_micro,
        values=noms_micros if noms_micros else ["(aucun micro détecté)"],
        state="readonly", width=72)
    combo_micro.pack(anchor="w", padx=22)
    if not noms_micros:
        combo_micro.configure(state="disabled")
        tk.Label(dlg, text="Aucun micro branché : seule la voix du patient sera transcrite.",
                 bg="#0d1117", fg="#d29922", font=("Segoe UI", 9)
                 ).pack(anchor="w", padx=22, pady=(4, 0))

    def valider():
        choix["sortie"] = var_sortie.get()
        choix["micro"] = var_micro.get() if noms_micros else ""
        choix["valide"] = True
        dlg.destroy()

    tk.Button(dlg, text="Démarrer", command=valider,
              font=("Segoe UI", 11, "bold")).pack(pady=20)
    dlg.bind("<Return>", lambda _e: valider())
    dlg.protocol("WM_DELETE_WINDOW", dlg.destroy)   # croix = ne pas demarrer
    dlg.mainloop()

    if not choix["valide"]:
        return None

    # Resolution avec repli sur le defaut si le peripherique a disparu.
    loopback = resoudre_loopback(choix["sortie"]) or loopback_defaut()
    if loopback is None:
        message_simple(
            "Sortie audio indisponible",
            "La sortie audio choisie n'est plus disponible.\n\nRelancez Écho.",
            genre="error")
        return None
    micro = resoudre_micro(choix["micro"]) if choix["micro"] else None

    # Persistance (best-effort) des choix.
    cfg["sortie"] = choix["sortie"]
    cfg["micro"] = choix["micro"]
    sauver_config(cfg)

    return loopback, micro, dossier_sauv


# ----------------------------- CAPTURE + VAD ---------------------------------

def capturer(source_factory, label):
    """Capture d'une source + segmentation par silence (webrtcvad).

    Robuste : tout probleme (peripherique absent, debranche en cours, erreur
    quelconque) coupe proprement CE flux avec un message simple, sans tuer
    l'autre flux ni l'application.
    """
    libelle = "patient" if label == "Patient" else "médecin"

    try:
        mic = source_factory()
    except Exception:
        mic = None
    if mic is None:
        display_queue.put(("AVIS",
            "La capture (%s) n'a pas pu démarrer. L'application reste utilisable." % libelle))
        return

    vad = webrtcvad.Vad(VAD_LEVEL)
    silence_frames_limit = SILENCE_MS // FRAME_MS
    min_speech_frames    = MIN_SPEECH_MS // FRAME_MS
    max_frames           = MAX_SEG_MS // FRAME_MS

    try:
        with mic.recorder(samplerate=SAMPLE_RATE, channels=1, blocksize=FRAME_SAMPLES) as rec:
            buffer = []
            silence_run = 0
            speaking = False

            while not stop_event.is_set():
                try:
                    data = rec.record(numframes=FRAME_SAMPLES)
                except Exception:
                    display_queue.put(("AVIS",
                        "La capture (%s) s'est interrompue (périphérique débranché ?). "
                        "L'autre source continue normalement." % libelle))
                    return
                mono = data[:, 0] if data.ndim > 1 else data
                gain = _get_gain(label)
                if gain != 1.0:
                    mono = np.clip(mono * gain, -1.0, 1.0, out=mono.copy())

                pcm16 = (np.clip(mono, -1.0, 1.0) * 32767).astype(np.int16)
                is_speech = vad.is_speech(pcm16.tobytes(), SAMPLE_RATE)
                _set_speaking(label, is_speech)

                if is_speech:
                    buffer.append(mono.astype(np.float32))
                    silence_run = 0
                    speaking = True
                elif speaking:
                    buffer.append(mono.astype(np.float32))
                    silence_run += 1

                fin_de_tour = speaking and silence_run >= silence_frames_limit
                trop_long   = len(buffer) >= max_frames

                if (fin_de_tour or trop_long) and len(buffer) >= min_speech_frames:
                    segment = np.concatenate(buffer)
                    segment_queue.put((label, segment))
                    buffer, silence_run, speaking = [], 0, False
                elif fin_de_tour:
                    buffer, silence_run, speaking = [], 0, False
    except Exception:
        display_queue.put(("AVIS",
            "Impossible de démarrer la capture (%s). Vérifiez le périphérique audio." % libelle))
    finally:
        _set_speaking(label, False)


# ----------------------------- MODELE / RESSOURCES ---------------------------

def ressource(rel):
    """Chemin d'une ressource, que l'appli soit en dev ou gelee (PyInstaller).
    En mode frozen, les --add-data sont extraits dans sys._MEIPASS."""
    base = getattr(sys, "_MEIPASS", os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base, rel)


def models_dir():
    """Dossier runtime des modeles telecharges : APPDATA/Echo/models/"""
    return os.path.join(dossier_config(), "models")


def whisper_ok():
    """Vérifie que le modèle Whisper est prêt (dossier + model.bin > 1 Go)."""
    d = os.path.join(models_dir(), MODELE_WHISPER_DIR)
    f = os.path.join(d, "model.bin")
    try:
        return os.path.isdir(d) and os.path.isfile(f) and os.path.getsize(f) > 1_000_000_000
    except OSError:
        return False


def qwen_ok():
    """Vérifie que le modèle Qwen est prêt (fichier .gguf > 1,5 Go)."""
    f = os.path.join(models_dir(), RESUME_GGUF)
    try:
        return os.path.isfile(f) and os.path.getsize(f) > 1_500_000_000
    except OSError:
        return False


def chemin_modele():
    """Résolution du chemin Whisper :
      - frozen + modele dans APPDATA : chemin APPDATA (installeur leger)
      - dev : dossier local du projet (pas de rupture du workflow dev)
      - fallback : MODEL_SIZE (telechargement HF)
    """
    # 1. APPDATA (runtime, frozen ou dev après téléchargement)
    if whisper_ok():
        return os.path.join(models_dir(), MODELE_WHISPER_DIR)
    # 2. Dev : dossier local dans le projet
    if not getattr(sys, "frozen", False):
        local = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                             MODELE_WHISPER_DIR)
        if os.path.isdir(local) and os.path.isfile(os.path.join(local, "model.bin")):
            return local
    # 3. Fallback : nom MODEL_SIZE → faster-whisper télécharge via HF
    return MODEL_SIZE


# ----------------------------- RÉSUMÉ (LLM local) ----------------------------
# Conception figee en Phase 2a : gabarit fixe, regle anti-invention, 2 exemples
# one-shot, temperature basse, et filets de securite cote code.

RESUME_GGUF        = "Qwen2.5-3B-Instruct-Q4_K_M.gguf"
RESUME_SEED        = 42
RESUME_TEMPERATURE = 0.1
RESUME_TOP_P       = 0.9
RESUME_MAX_TOKENS  = 450

# Avertissement non negociable, ajoute PAR LE CODE (toujours exact).
ENTETE_RESUME = "RÉSUMÉ (généré automatiquement — à relire et corriger)"

RESUME_TITRES = [
    "Motif :",
    "Observations / points clés :",
    "Traitements et prescriptions évoqués :",
    "Suivi et recommandations :",
]

# Formules de politesse / cloture, jamais cliniques -> retirees cote code.
RESUME_POLITESSE = re.compile(
    r"bonne (journée|soirée|saison|continuation|route)|"
    r"prendre soin|prenez soin|au revoir|à bient[oô]t|portez-vous bien|"
    r"bonne fin de", re.IGNORECASE)

RESUME_SYSTEM = (
    "Tu es un assistant médical. Tu rédiges le compte-rendu d'une consultation "
    "à partir de sa transcription.\n\n"
    "RÈGLES IMPÉRATIVES :\n"
    "- Utilise UNIQUEMENT les informations présentes dans la transcription. "
    "N'invente aucun diagnostic, médicament, posologie ni chiffre non énoncé.\n"
    "- Rédige de façon impersonnelle, à la troisième personne. Ne t'adresse JAMAIS "
    "au patient (n'emploie jamais « vous »).\n"
    "- Respecte les négations : un symptôme nié par le patient ne doit JAMAIS "
    "apparaître comme présent ; écris-le sous forme négative claire "
    "(ex. « Pas de toux, pas de fièvre »).\n"
    "- Chaque section est une liste de puces courtes commençant par « - », même "
    "s'il n'y a qu'un seul élément.\n"
    "- Ignore les salutations et formules de politesse (bonjour, au revoir, merci, "
    "bonne journée, prenez soin de vous, etc.) : elles n'apparaissent pas dans le "
    "compte-rendu.\n"
    "- Si, et seulement si, une section ne contient réellement aucune information "
    "dans la transcription, écris sur une seule ligne, sans puce : Non précisé. "
    "Ne déduis rien, n'extrapole aucune recommandation non formulée.\n"
    "- Aucune phrase d'introduction ni de conclusion, aucun autre titre que les "
    "quatre imposés.\n\n"
    "CONTENU DE CHAQUE SECTION :\n"
    "- Motif : la raison de la consultation.\n"
    "- Observations / points clés : symptômes et plaintes décrits par le patient "
    "ET résultats de l'examen clinique (auscultation, tension, poids, etc.).\n"
    "- Traitements et prescriptions évoqués : médicaments, examens prescrits, "
    "orientations vers un spécialiste.\n"
    "- Suivi et recommandations : prochain rendez-vous, consignes de surveillance, "
    "conduite à tenir en cas d'aggravation.\n\n"
    "FORMAT — réponds EXACTEMENT avec ces quatre titres, dans cet ordre, suivis "
    "de deux points :\n\n"
    "Motif :\n"
    "Observations / points clés :\n"
    "Traitements et prescriptions évoqués :\n"
    "Suivi et recommandations :"
)

RESUME_ONESHOT1_USER = (
    "Transcription de la consultation :\n\n"
    "[09:00:00] Médecin : Bonjour, qu'est-ce qui vous amène ?\n"
    "[09:00:05] Patient : J'ai mal à la gorge depuis deux jours, avec un peu de fièvre.\n"
    "[09:00:12] Médecin : Pas de toux, pas de gêne pour respirer ?\n"
    "[09:00:16] Patient : Non, pas de toux, je respire bien.\n"
    "[09:00:22] Médecin : La gorge est rouge, les ganglions du cou sont un peu gonflés. "
    "Température 38,2.\n"
    "[09:00:40] Médecin : C'est une angine probablement virale. Paracétamol 1 gramme "
    "si douleur ou fièvre, trois fois par jour maximum.\n"
    "[09:00:55] Médecin : Reposez-vous, buvez beaucoup. Si la fièvre dépasse trois "
    "jours, revenez consulter.\n\n"
    "Rédige le compte-rendu selon le format imposé."
)
RESUME_ONESHOT1_ASSISTANT = (
    "Motif :\n"
    "- Mal de gorge depuis deux jours avec fièvre.\n\n"
    "Observations / points clés :\n"
    "- Gorge rouge, ganglions cervicaux légèrement gonflés.\n"
    "- Température à 38,2 °C.\n"
    "- Pas de toux, pas de gêne respiratoire.\n\n"
    "Traitements et prescriptions évoqués :\n"
    "- Paracétamol 1 g si douleur ou fièvre, trois fois par jour maximum.\n\n"
    "Suivi et recommandations :\n"
    "- Repos et hydratation abondante.\n"
    "- Reconsulter si la fièvre dépasse trois jours."
)
RESUME_ONESHOT2_USER = (
    "Transcription de la consultation :\n\n"
    "[10:00:00] Médecin : Bonjour, je vous écoute.\n"
    "[10:00:04] Patient : Je voulais juste savoir si je peux prendre du paracétamol "
    "avec mon traitement habituel.\n"
    "[10:00:10] Médecin : Oui, c'est compatible, aucun problème.\n"
    "[10:00:15] Patient : Parfait, merci, c'était seulement ça.\n"
    "[10:00:18] Médecin : Très bien, bonne journée.\n\n"
    "Rédige le compte-rendu selon le format imposé."
)
RESUME_ONESHOT2_ASSISTANT = (
    "Motif :\n"
    "- Question sur la compatibilité du paracétamol avec le traitement habituel.\n\n"
    "Observations / points clés :\n"
    "Non précisé\n\n"
    "Traitements et prescriptions évoqués :\n"
    "Non précisé\n\n"
    "Suivi et recommandations :\n"
    "Non précisé"
)

_resume_llm = None
_resume_lock = threading.Lock()


def chemin_modele_resume():
    """Résolution du chemin Qwen (.gguf) :
      - APPDATA (runtime, frozen ou dev après téléchargement)
      - Dev : bench/models/ local (workflow dev intact)
    """
    # 1. APPDATA
    if qwen_ok():
        return os.path.join(models_dir(), RESUME_GGUF)
    # 2. Dev local
    if not getattr(sys, "frozen", False):
        dev = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           "bench", "models", RESUME_GGUF)
        if os.path.isfile(dev):
            return dev
    raise FileNotFoundError(RESUME_GGUF)


def charger_modele_resume():
    """Charge le modele UNE fois et le garde en memoire pour toute la session
    (rechargement evite si la generation est relancee). Thread-safe."""
    global _resume_llm
    with _resume_lock:
        if _resume_llm is None:
            from llama_cpp import Llama
            _resume_llm = Llama(
                model_path=chemin_modele_resume(), n_ctx=8192,
                n_threads=multiprocessing.cpu_count(),
                seed=RESUME_SEED, verbose=False)
        return _resume_llm


def _normaliser_resume(corps):
    """Garantit chaque titre sur sa propre ligne, retire les puces de politesse
    et les lignes vides multiples."""
    for t in RESUME_TITRES:
        corps = corps.replace(t, "\n" + t + "\n")
    lignes = []
    for l in corps.splitlines():
        l = l.rstrip()
        if l.lstrip().startswith("-") and RESUME_POLITESSE.search(l):
            continue
        if l == "" and (not lignes or lignes[-1] == ""):
            continue
        lignes.append(l)
    return "\n".join(lignes).strip()


def generer_resume(llm, transcript):
    """Genere le resume structure (en-tete fixe + 4 sections normalisees)."""
    out = llm.create_chat_completion(
        messages=[
            {"role": "system", "content": RESUME_SYSTEM},
            {"role": "user", "content": RESUME_ONESHOT1_USER},
            {"role": "assistant", "content": RESUME_ONESHOT1_ASSISTANT},
            {"role": "user", "content": RESUME_ONESHOT2_USER},
            {"role": "assistant", "content": RESUME_ONESHOT2_ASSISTANT},
            {"role": "user",
             "content": "Transcription de la consultation :\n\n" + transcript
                        + "\n\nRédige le compte-rendu selon le format imposé."},
        ],
        temperature=RESUME_TEMPERATURE,
        top_p=RESUME_TOP_P,
        max_tokens=RESUME_MAX_TOKENS,
        seed=RESUME_SEED,
    )
    corps = _normaliser_resume(out["choices"][0]["message"]["content"].strip())
    return ENTETE_RESUME + "\n\n" + corps


# ----------------------------- PARAMÈTRES WHISPER ----------------------------

# ~80 termes médicaux / médicaments pour guider le tokenizer Whisper.
# Ajouter ici toute orthographe que le modèle small manque régulièrement.
WHISPER_INITIAL_PROMPT = (
    # Médicaments courants
    "Doliprane, paracétamol, ibuprofène, amoxicilline, Augmentin, Efferalgan, "
    "Spasfon, Ventoline, Smecta, Toplexil, Pivalone, Xyzall, Voltaren, Dafalgan, "
    "metformine, amlodipine, bisoprolol, ramipril, atorvastatine, lévothyroxine, "
    "oméprazole, pantoprazole, Inexium, Kardégic, Plavix, Eliquis, Xarelto, "
    "corticoïdes, antibiotique, antihistaminique, bronchodilatateur, "
    # Termes cliniques
    "auscultation, palpations, palpitations, dyspnée, essoufflement, "
    "tachycardie, bradycardie, hypertension, hypotension, saturation, "
    "fièvre, frissons, nausées, vomissements, diarrhée, constipation, "
    "céphalées, vertiges, syncope, œdème, cicatrisation, inflammation, "
    "infection, allergie, diabète, "
    # Examens et actes
    "bilan sanguin, NFS, ferritine, CRP, TSH, échographie, radio, scanner, IRM, "
    "ordonnance, renouvellement, arrêt de travail, certificat médical, "
    "vaccination, rappel, consultation, tension artérielle."
)

# Dictionnaire de corrections post-transcription.
# Clés : erreurs connues du modèle small en français médical (insensibles à la casse).
# Valeurs : orthographe correcte.
# Facile à étendre : ajouter une ligne "erreur": "correction".
_CORRECTIONS_RAW = {
    "dolifren":       "Doliprane",
    "dolipren":       "Doliprane",
    "doliprant":      "Doliprane",
    "ibuprofène":     "ibuprofène",   # normalise les variantes sans accent
    "ibuprofen":      "ibuprofène",
    "amoxiciline":    "amoxicilline",
    "la fière":       "la fièvre",
    "une fière":      "une fièvre",
    "de fière":       "de fièvre",
    "ventoline":      "Ventoline",
    "levothyroxine":  "lévothyroxine",
    "metformin":      "metformine",
    "oeudeme":        "œdème",
    "oedème":         "œdème",
    "dynspnée":       "dyspnée",
    "dyspné":         "dyspnée",
    "palitations":    "palpitations",
    "palpation":      "palpitations",  # ambiguïté fréquente à l'oral
    "tension":        "tension",       # placeholder — ne remplace pas, sert de doc
}

# On compile une table { minuscules: correction } pour la recherche rapide.
_CORRECTIONS = {k.lower(): v for k, v in _CORRECTIONS_RAW.items()
                if k.lower() != v.lower()}  # ignore les no-ops


def corriger_transcription(texte):
    """Applique les corrections post-transcription (orthographe médicale).
    Remplacement mot-à-mot insensible à la casse, respecte la ponctuation."""
    if not texte or not _CORRECTIONS:
        return texte
    # Tokenise en préservant les séparateurs, remplace mot par mot.
    tokens = re.split(r'(\W+)', texte)
    result = []
    i = 0
    while i < len(tokens):
        t = tokens[i]
        tl = t.lower()
        # Essaie d'abord les bi-grammes (ex. "la fièvre").
        if i + 2 < len(tokens):
            bigram = (t + tokens[i + 1] + tokens[i + 2]).lower()
            if bigram in _CORRECTIONS:
                result.append(_CORRECTIONS[bigram])
                i += 3
                continue
        if tl in _CORRECTIONS:
            result.append(_CORRECTIONS[tl])
        else:
            result.append(t)
        i += 1
    return "".join(result)


# ----------------------------- TRANSCRIPTION ---------------------------------

def transcrire():
    """Worker unique : consomme les segments et les transcrit (modele non concurrent)."""
    source = chemin_modele()
    display_queue.put(("INFO", "Chargement du modele (" + source + ")..."))
    model = WhisperModel(source, device=DEVICE, compute_type=COMPUTE_TYPE)
    display_queue.put(("INFO", "Pret. La transcription demarre."))

    while not stop_event.is_set():
        try:
            label, audio = segment_queue.get(timeout=0.5)
        except queue.Empty:
            continue

        segments, _ = model.transcribe(
            audio, language=LANGUAGE, vad_filter=True,
            beam_size=3,
            condition_on_previous_text=True,
            initial_prompt=WHISPER_INITIAL_PROMPT,
        )
        texte = corriger_transcription("".join(s.text for s in segments).strip())
        if texte:
            display_queue.put((label, texte))


# ----------------------------- EXPORT DOCUMENT ------------------------------

# Libelle accentue ecrit dans le fichier (l'interne reste sans accent).
LOCUTEUR_FICHIER = {"Medecin": "Médecin", "Patient": "Patient"}

# Couleurs pour la section transcription.
_BLEU_PATIENT = "2E74B5"   # bleu sobre (hex sans #)
_GRIS_TITRE   = "404040"   # gris foncé pour Heading TRANSCRIPTION

def _pt(points):
    """Convertit des points en unités Emu (utilisées par python-docx)."""
    from docx.shared import Pt
    return Pt(points)

def _cm(centimetres):
    from docx.shared import Cm
    return Cm(centimetres)

def _rgb(hex6):
    """Retourne un objet RGBColor depuis une chaine hex sans #."""
    from docx.shared import RGBColor
    r, g, b = int(hex6[0:2], 16), int(hex6[2:4], 16), int(hex6[4:6], 16)
    return RGBColor(r, g, b)

def _forcer_style_heading(style, taille, hex_couleur, gras=True, police="Arial"):
    """Surcharge un style Heading : police Arial, couleur imposee (non bleue)."""
    from docx.shared import Pt, RGBColor
    style.font.name = police
    style.font.size = Pt(taille)
    style.font.bold = gras
    style.font.color.rgb = _rgb(hex_couleur)

def _set_interligne(para, multiple):
    """Fixe l'interligne d'un paragraphe (multiple, ex 1.2)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_LINE_SPACING
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = multiple

def _ajouter_run(para, texte, gras=False, italique=False,
                 couleur=None, taille=12, police="Arial"):
    """Ajoute un run formaté dans un paragraphe."""
    from docx.shared import Pt
    r = para.add_run(texte)
    r.font.name = police
    r.font.size = Pt(taille)
    r.bold = gras
    r.italic = italique
    if couleur:
        r.font.color.rgb = _rgb(couleur)
    return r

def _ecrire_docx(chemin, infos, now, resume, entries, annexes=None):
    """Construit et ecrit le .docx selon la structure demandee."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # --- Marges A4 2,5 cm ---
    section = doc.sections[0]
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = section.right_margin = Cm(2.5)
        section.top_margin  = section.bottom_margin = Cm(2.5)
    section = doc.sections[0]   # reference gardee pour la largeur images

    # --- Surcharger Heading 1 (noir) et Heading 2 (noir) ---
    _forcer_style_heading(doc.styles["Heading 1"], 14, "000000")
    _forcer_style_heading(doc.styles["Heading 2"], 13, "000000")

    # 1. Titre principal centré.
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_interligne(titre, 1.2)
    run_titre = titre.add_run("Compte-rendu de consultation")
    run_titre.font.name = "Arial"
    run_titre.font.size = Pt(18)
    run_titre.bold = True
    run_titre.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()   # espace

    # 2. Tableau patient 2 colonnes.
    date_str  = now.strftime("%d/%m/%Y")
    heure_str = now.strftime("%Hh%M")
    lignes_patient = [
        ("Nom",       infos["nom"]),
        ("Prénom",    infos["prenom"]),
        ("Né(e) le",  infos["naissance"]),
        ("Date",      "%s à %s" % (date_str, heure_str)),
        ("Motif",     infos.get("motif") or "—"),
    ]
    tbl = doc.add_table(rows=len(lignes_patient), cols=2)
    tbl.style = "Table Grid"
    for row_idx, (lbl, val) in enumerate(lignes_patient):
        c0, c1 = tbl.rows[row_idx].cells
        # Largeur colonnes
        c0.width = Cm(4.5)
        c1.width = Cm(12.5)
        # Étiquette en gras
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(lbl)
        r0.bold = True
        r0.font.name = "Arial"
        r0.font.size = Pt(12)
        # Valeur
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.name = "Arial"
        r1.font.size = Pt(12)
        # Couleur de fond légère pour l'étiquette (gris très clair)
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OE
        shd = _OE("w:shd")
        shd.set(_qn("w:val"), "clear")
        shd.set(_qn("w:color"), "auto")
        shd.set(_qn("w:fill"), "F2F2F2")
        c0.paragraphs[0]._p.get_or_add_pPr().append(shd)
    doc.add_paragraph()   # espace après tableau

    # 3. Séparateur horizontal (paragraphe avec bordure bottom).
    sep_para = doc.add_paragraph()
    sep_pPr = sep_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    sep_pPr.append(pBdr)
    doc.add_paragraph()

    # 4. Section RÉSUMÉ (si présent).
    if resume:
        h1_res = doc.add_heading("Résumé de la consultation", level=1)
        _set_interligne(h1_res, 1.2)

        # Bandeau avertissement en italique grisé.
        avert = doc.add_paragraph()
        _set_interligne(avert, 1.2)
        ra = avert.add_run(
            "⚠ Résumé généré automatiquement — à relire et corriger par le médecin.")
        ra.italic = True
        ra.font.name = "Arial"
        ra.font.size = Pt(11)
        ra.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        doc.add_paragraph()

        # Découper le résumé en sections.
        TITRES_RESUME = [
            "Motif :",
            "Observations / points clés :",
            "Traitements et prescriptions évoqués :",
            "Suivi et recommandations :",
        ]
        lignes_res = resume.splitlines()
        # Sauter la première ligne (ENTETE_RESUME déjà dans bandeau).
        debut = 0
        for idx, l in enumerate(lignes_res):
            if l.strip() and not l.startswith("RÉSUMÉ"):
                debut = idx
                break

        titre_courant = None
        puces_courantes = []

        def _flush_section():
            if titre_courant is None:
                return
            h2 = doc.add_heading(titre_courant.rstrip(":"), level=2)
            _set_interligne(h2, 1.2)
            if puces_courantes == ["Non précisé"] or not puces_courantes:
                p = doc.add_paragraph("Non précisé",
                                      style="List Bullet" if False else "Normal")
                _set_interligne(p, 1.2)
                p.runs[0].font.name = "Arial"
                p.runs[0].font.size = Pt(12)
                p.runs[0].italic = True
            else:
                for puce in puces_courantes:
                    p = doc.add_paragraph(style="List Bullet")
                    _set_interligne(p, 1.2)
                    _ajouter_run(p, puce, taille=12)

        for ligne in lignes_res[debut:]:
            ligne = ligne.strip()
            if not ligne:
                continue
            matched = next((t for t in TITRES_RESUME if ligne.startswith(t)), None)
            if matched:
                _flush_section()
                titre_courant = matched
                puces_courantes = []
            elif ligne == "Non précisé":
                puces_courantes = ["Non précisé"]
            elif ligne.startswith("- "):
                puces_courantes.append(ligne[2:].strip())
            elif ligne.startswith("-"):
                puces_courantes.append(ligne[1:].strip())
        _flush_section()
        doc.add_paragraph()

    # 5. Saut de page.
    doc.add_page_break()

    # 6. Section TRANSCRIPTION INTÉGRALE.
    h1_tr = doc.add_heading("Transcription intégrale", level=1)
    _set_interligne(h1_tr, 1.2)
    h1_tr.runs[0].font.color.rgb = _rgb(_GRIS_TITRE)

    for horodatage, locuteur, texte in entries:
        loc_label = LOCUTEUR_FICHIER.get(locuteur, locuteur)
        p = doc.add_paragraph()
        _set_interligne(p, 1.15)
        est_medecin = (locuteur == "Medecin")
        # Préfixe locuteur + heure.
        prefixe = "%s (%s) : " % (loc_label, horodatage)
        r_pre = p.add_run(prefixe)
        r_pre.font.name = "Arial"
        r_pre.font.size = Pt(11)
        r_pre.bold = est_medecin
        if not est_medecin:
            r_pre.font.color.rgb = _rgb(_BLEU_PATIENT)
        # Texte transcrit.
        r_txt = p.add_run(texte)
        r_txt.font.name = "Arial"
        r_txt.font.size = Pt(11)
        if not est_medecin:
            r_txt.font.color.rgb = _rgb(_BLEU_PATIENT)

    # 7. Section DOCUMENTS ANNEXES (si au moins un document selectionne).
    if annexes:
        _inserer_annexes_docx(doc, annexes, section)

    doc.save(chemin)


def _pdf_vers_image_bytes(chemin_pdf):
    """Convertit la 1re page d'un PDF en PNG (bytes) via PyMuPDF.
    Renvoie les bytes PNG, ou None si PyMuPDF indisponible / erreur."""
    try:
        import fitz
        doc_pdf = fitz.open(chemin_pdf)
        page = doc_pdf[0]
        pix = page.get_pixmap(dpi=150)
        return pix.tobytes("png")
    except Exception:
        return None


def _inserer_annexes_docx(doc, annexes, page_section):
    """Ajoute la section DOCUMENTS ANNEXES apres la transcription.
    annexes : liste de {"label": str, "fichier": str|None}
    Chaque document est independant — un echec n'arrete pas les autres.
    """
    import io
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    # Séparateur horizontal.
    sep_para = doc.add_paragraph()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    sep_para._p.get_or_add_pPr().append(pBdr)

    h1_ann = doc.add_heading("Documents annexes", level=1)
    _set_interligne(h1_ann, 1.2)
    h1_ann.runs[0].font.color.rgb = _rgb(_GRIS_TITRE)

    # Largeur utile de la page (pour les images).
    larg_page = (page_section.page_width
                 - page_section.left_margin
                 - page_section.right_margin)

    for doc_info in annexes:
        label   = doc_info.get("label", "Document")
        fichier = doc_info.get("fichier")   # peut être None

        h2 = doc.add_heading(label, level=2)
        _set_interligne(h2, 1.2)

        if not fichier:
            # Case cochée mais aucun fichier joint.
            p = doc.add_paragraph()
            _set_interligne(p, 1.2)
            _ajouter_run(p, "Document remis au patient lors de la consultation.",
                         italique=True, couleur="808080", taille=11)
            continue

        ext = os.path.splitext(fichier)[1].lower()

        # --- Image directe (JPG / PNG) ---
        if ext in (".jpg", ".jpeg", ".png"):
            try:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(fichier, width=larg_page)
            except Exception:
                p = doc.add_paragraph()
                _ajouter_run(p, "Image jointe — impossible d'incorporer le fichier.",
                             italique=True, couleur="808080", taille=11)
            continue

        # --- PDF : tentative de conversion via PyMuPDF ---
        if ext == ".pdf":
            img_bytes = _pdf_vers_image_bytes(fichier)
            if img_bytes:
                try:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(
                        io.BytesIO(img_bytes), width=larg_page)
                    img_bytes = img_bytes   # marqueur succès
                except Exception:
                    img_bytes = None
            if not img_bytes:
                p = doc.add_paragraph()
                _ajouter_run(
                    p,
                    "Document PDF joint lors de la consultation — "
                    "conserver séparément.",
                    italique=True, couleur="808080", taille=11)
            continue

        # --- Autre format ---
        p = doc.add_paragraph()
        _ajouter_run(p, "Fichier joint : " + os.path.basename(fichier),
                     italique=True, couleur="808080", taille=11)


def _ecrire_txt_secours(chemin, infos, now, resume, entries):
    """Fallback .txt si python-docx echoue (la transcription ne doit jamais etre perdue)."""
    sep = "=" * 60
    with open(chemin, "w", encoding="utf-8") as f:
        f.write(sep + "\n")
        f.write("  CONSULTATION — %s %s\n" % (infos["nom"].upper(), infos["prenom"]))
        f.write("  Né(e) le : %s\n" % infos["naissance"])
        f.write("  Date : %s à %s\n" % (now.strftime("%d/%m/%Y"), now.strftime("%Hh%M")))
        f.write("  Motif : %s\n" % (infos.get("motif") or "—"))
        f.write(sep + "\n\n")
        if resume:
            f.write(resume.strip() + "\n\n")
            f.write(sep + "\n  TRANSCRIPTION INTÉGRALE\n" + sep + "\n\n")
        for horodatage, locuteur, texte in entries:
            f.write("[%s] %s : %s\n" % (
                horodatage, LOCUTEUR_FICHIER.get(locuteur, locuteur), texte))


# ----------------------------- INTERFACE -------------------------------------


def nettoyer_nom_fichier(s):
    """Rend `s` utilisable dans un nom de fichier Windows : sans accents ni
    caracteres interdits. (Les accents restent intacts dans l'en-tete du fichier.)"""
    s = unicodedata.normalize("NFKD", s)
    s = "".join(c for c in s if not unicodedata.combining(c))
    s = re.sub(r'[<>:"/\\|?*]', "", s)          # caracteres interdits Windows
    s = re.sub(r"[\x00-\x1f]", "", s)            # caracteres de controle
    s = re.sub(r"\s+", "-", s.strip())
    return s or "consultation"


class Overlay:
    COULEURS = {
        "Medecin": "#7ee787",
        "Patient": "#79c0ff",
        "INFO":    "#8b949e",
        "AVIS":    "#d29922",
        "ERREUR":  "#ff7b72",
    }

    def __init__(self, root, dossier_sauvegarde=None, banniere=None):
        self.root = root
        # Transcript complet accumule au fil de l'eau : (horodatage, locuteur, texte).
        self.entries = []
        # Dossier de sauvegarde memorise (ouverture par defaut de la boite fichier).
        self.dossier_sauvegarde = dossier_sauvegarde

        root.title("Écho — transcription consultation")
        root.geometry("520x640+40+40")
        root.attributes("-topmost", True)
        root.configure(bg="#0d1117")

        # Bandeau doux (ex : aucun micro detecte), affiche en haut si fourni.
        if banniere:
            tk.Label(root, text=banniere, bg="#3a2d10", fg="#f0c674",
                     font=("Segoe UI", 10), wraplength=496, justify="left",
                     padx=12, pady=8).pack(side="top", fill="x")

        # Barre du bas avec le bouton Quitter.
        barre = tk.Frame(root, bg="#161b22")
        barre.pack(side="bottom", fill="x")
        tk.Button(barre, text="Quitter", command=self.on_quit,
                  font=("Segoe UI", 10, "bold")).pack(side="right", padx=10, pady=8)

        self.text = tk.Text(
            root, bg="#0d1117", fg="#e6edf3", font=("Segoe UI", 13),
            wrap="word", padx=14, pady=14, borderwidth=0, spacing1=4, spacing3=6,
        )
        self.text.pack(side="top", fill="both", expand=True)
        for nom, couleur in self.COULEURS.items():
            self.text.tag_configure(nom, foreground=couleur)
        self.text.configure(state="disabled")

        # La croix de fermeture passe par le meme chemin que le bouton Quitter.
        root.protocol("WM_DELETE_WINDOW", self.on_quit)

        self.root.after(100, self._vider_file)

    def _vider_file(self):
        while True:
            try:
                label, texte = display_queue.get_nowait()
            except queue.Empty:
                break
            self.text.configure(state="normal")
            if label in ("INFO", "AVIS", "ERREUR"):
                self.text.insert("end", "- " + texte + "\n", label)
            else:
                # Entree de transcription : on l'affiche ET on l'accumule.
                horodatage = datetime.datetime.now().strftime("%H:%M:%S")
                self.entries.append((horodatage, label, texte))
                self.text.insert("end", label + " : ", label)
                self.text.insert("end", texte + "\n")
            self.text.see("end")
            self.text.configure(state="disabled")
        if not stop_event.is_set():
            self.root.after(100, self._vider_file)

    # ------------------------- Fermeture / enregistrement --------------------

    def on_quit(self):
        """Bouton Quitter ET croix de fermeture. Tout sur le thread principal."""
        rep = messagebox.askyesnocancel(
            "Quitter",
            "Enregistrer la transcription de cette consultation ?",
            parent=self.root,
        )
        if rep is None:
            return                      # Annuler -> on revient a l'appli
        if rep:                         # Oui -> formulaire patient, resume, fichier
            infos = self._formulaire_patient()
            if infos is None:
                return                  # formulaire annule -> on revient a l'appli

            resume = None
            if self.entries:            # rien a resumer si aucune parole captee
                resume = self._generer_resume_avec_progres(self._verbatim_texte())
                if resume is not None:
                    resume = self._relire_resume(resume)
                    if resume is None:
                        return          # relecture annulee -> on revient a l'appli
                else:
                    # Echec modele/generation : pas d'erreur technique, on continue.
                    message_simple(
                        "Résumé indisponible",
                        "Le résumé automatique n'a pas pu être généré.\n\n"
                        "La transcription mot à mot sera enregistrée normalement.")

            # Fenêtre documents annexes (après relecture, avant sauvegarde).
            annexes = self._choisir_annexes()

            if not self._enregistrer(infos, resume, annexes):
                return                  # boite de fichier annulee -> on revient a l'appli
        # Non, ou enregistrement effectue -> fermeture effective.
        self._fermer()

    def _fermer(self):
        stop_event.set()
        self.root.destroy()

    def _verbatim_texte(self):
        """Transcription mot a mot accumulee, au format horodate (pour le modele
        et pour le fichier)."""
        return "\n".join(
            "[%s] %s : %s" % (h, LOCUTEUR_FICHIER.get(loc, loc), t)
            for h, loc, t in self.entries)

    def _generer_resume_avec_progres(self, transcript):
        """Genere le resume dans un THREAD (UI non gelee), avec indicateur en
        deux temps (chargement puis redaction). Renvoie le resume ou None si
        echec. Le modele reste en memoire pour la session."""
        dlg = tk.Toplevel(self.root)
        dlg.title("Résumé")
        dlg.configure(bg="#0d1117")
        dlg.attributes("-topmost", True)
        dlg.resizable(False, False)
        dlg.grab_set()
        dlg.protocol("WM_DELETE_WINDOW", lambda: None)  # pas de fermeture en cours

        lbl = tk.Label(dlg, text="Chargement du modèle de résumé...",
                       bg="#0d1117", fg="#e6edf3", font=("Segoe UI", 12))
        lbl.pack(padx=36, pady=(26, 6))
        tk.Label(dlg, text="Cela peut prendre une à deux minutes selon l'ordinateur.",
                 bg="#0d1117", fg="#8b949e", font=("Segoe UI", 9)).pack(padx=36, pady=(0, 26))

        etat = {"stage": "load", "resume": None, "done": False}

        def worker():
            try:
                llm = charger_modele_resume()
                etat["stage"] = "gen"
                etat["resume"] = generer_resume(llm, transcript)
            except Exception:
                etat["resume"] = None
            etat["done"] = True

        threading.Thread(target=worker, daemon=True).start()

        def tick():
            if etat["done"]:
                try:
                    dlg.destroy()
                except Exception:
                    pass
                return
            lbl.config(text=("Chargement du modèle de résumé..."
                             if etat["stage"] == "load"
                             else "Rédaction du résumé en cours..."))
            dlg.after(200, tick)

        dlg.after(200, tick)
        self.root.wait_window(dlg)
        return etat["resume"]

    def _relire_resume(self, resume):
        """Brouillon editable : le medecin relit/corrige avant enregistrement.
        Renvoie le texte valide, ou None si annulation.

        Structure (de haut en bas, empile avant la zone texte) :
          1. Bandeau avertissement  — fixe en haut  (pack side=TOP)
          2. Frame boutons          — fixe en bas   (pack side=BOTTOM avant le texte)
          3. Zone texte + scrollbar — remplit le reste (pack side=TOP, expand=True)
        Le pack BOTTOM avant TOP garantit que la barre de boutons ne est jamais
        chassee hors ecran quand le texte est grand.
        """
        dlg = tk.Toplevel(self.root)
        dlg.title("Résumé — à relire et corriger")
        dlg.configure(bg="#0d1117")
        dlg.attributes("-topmost", True)
        dlg.geometry("660x560")
        dlg.minsize(520, 480)           # boutons toujours visibles a l'ouverture
        dlg.grab_set()

        # 1 — Bandeau fixe en haut.
        tk.Label(dlg,
                 text="Relisez et corrigez le résumé avant l'enregistrement. "
                      "La transcription mot à mot sera conservée dans le fichier.",
                 bg="#3a2d10", fg="#f0c674", font=("Segoe UI", 10),
                 wraplength=630, justify="left").pack(
                     side="top", fill="x", padx=0, pady=0, ipadx=12, ipady=8)

        res = {"valide": None}

        def valider():
            res["valide"] = txt.get("1.0", "end").strip()
            dlg.destroy()

        def annuler():
            dlg.destroy()               # res["valide"] reste None

        # 2 — Barre de boutons fixe en bas (packer AVANT la zone texte).
        barre = tk.Frame(dlg, bg="#161b22")
        barre.pack(side="bottom", fill="x")
        tk.Button(
            barre, text="Valider et enregistrer", command=valider,
            font=("Segoe UI", 11, "bold"), width=22,
        ).pack(side="right", padx=14, pady=10)
        tk.Button(
            barre, text="Annuler", command=annuler,
            font=("Segoe UI", 11), width=18,
        ).pack(side="right", padx=(0, 6), pady=10)

        # 3 — Zone texte + scrollbar (prend tout l'espace restant).
        frame_txt = tk.Frame(dlg, bg="#0d1117")
        frame_txt.pack(side="top", fill="both", expand=True)

        sb = tk.Scrollbar(frame_txt)
        sb.pack(side="right", fill="y")

        txt = tk.Text(frame_txt, bg="#0d1117", fg="#e6edf3",
                      font=("Segoe UI", 11), wrap="word",
                      padx=12, pady=12, insertbackground="#e6edf3",
                      yscrollcommand=sb.set)
        txt.pack(side="left", fill="both", expand=True)
        sb.config(command=txt.yview)
        txt.insert("1.0", resume)

        dlg.protocol("WM_DELETE_WINDOW", annuler)
        self.root.wait_window(dlg)
        return res["valide"]

    def _formulaire_patient(self):
        """Toplevel modal de saisie des infos patient. Renvoie un dict ou None
        si l'utilisateur annule. Nom/Prenom/Date de naissance obligatoires."""
        dlg = tk.Toplevel(self.root)
        dlg.title("Informations patient")
        dlg.configure(bg="#0d1117")
        dlg.attributes("-topmost", True)
        dlg.resizable(False, False)
        dlg.grab_set()                  # modal

        resultat = {"data": None}
        champs = {}

        def champ(libelle):
            tk.Label(dlg, text=libelle, bg="#0d1117", fg="#e6edf3",
                     font=("Segoe UI", 10)).pack(anchor="w", padx=18, pady=(10, 0))
            e = tk.Entry(dlg, width=44, font=("Segoe UI", 11))
            e.pack(padx=18, pady=(0, 2))
            return e

        champs["nom"]       = champ("Nom *")
        champs["prenom"]    = champ("Prénom *")
        champs["naissance"] = champ("Date de naissance (JJ/MM/AAAA) *")
        champs["motif"]     = champ("Motif de consultation (facultatif)")

        err = tk.Label(dlg, text="", bg="#0d1117", fg="#ff7b72", font=("Segoe UI", 9))
        err.pack(padx=18, pady=(6, 0))

        def valider():
            nom       = champs["nom"].get().strip()
            prenom    = champs["prenom"].get().strip()
            naissance = champs["naissance"].get().strip()
            motif     = champs["motif"].get().strip()
            if not nom:
                err.config(text="Le champ Nom est obligatoire.")
                champs["nom"].focus_set()
                return
            if not prenom:
                err.config(text="Le champ Prénom est obligatoire.")
                champs["prenom"].focus_set()
                return
            if not naissance:
                err.config(text="La date de naissance est obligatoire.")
                champs["naissance"].focus_set()
                return
            resultat["data"] = {"nom": nom, "prenom": prenom,
                                "naissance": naissance, "motif": motif}
            dlg.destroy()

        def annuler():
            dlg.destroy()               # resultat reste None

        btns = tk.Frame(dlg, bg="#0d1117")
        btns.pack(pady=14)
        tk.Button(btns, text="Valider", command=valider,
                  font=("Segoe UI", 10, "bold")).pack(side="left", padx=6)
        tk.Button(btns, text="Annuler", command=annuler).pack(side="left", padx=6)

        dlg.protocol("WM_DELETE_WINDOW", annuler)
        champs["nom"].focus_set()
        self.root.wait_window(dlg)
        return resultat["data"]

    def _choisir_annexes(self):
        """Fenêtre 'Documents annexes' (après relecture résumé, avant sauvegarde).
        Renvoie une liste de {"label": str, "fichier": str|None} (jamais None).
        Fermeture sans choix = liste vide (équivalent 'Passer')."""
        DOCS = [
            ("ordonnance",    "Ordonnance"),
            ("arret",         "Arrêt de travail"),
            ("autre",         "Autre document"),
        ]
        resultat = {"annexes": [], "valide": False}

        dlg = tk.Toplevel(self.root)
        dlg.title("Documents créés lors de cette consultation")
        dlg.configure(bg="#0d1117")
        dlg.attributes("-topmost", True)
        dlg.geometry("620x480")
        dlg.minsize(540, 380)
        dlg.grab_set()

        # --- En-tête ---
        tk.Label(dlg, text="Documents créés lors de cette consultation",
                 bg="#0d1117", fg="#e6edf3",
                 font=("Segoe UI", 13, "bold")).pack(anchor="w", padx=22, pady=(22, 2))
        tk.Label(dlg,
                 text="Sélectionnez les documents à inclure dans le compte-rendu.\n"
                      "Vous pourrez joindre un fichier (photo ou scan) pour chacun.",
                 bg="#0d1117", fg="#8b949e", font=("Segoe UI", 10),
                 justify="left").pack(anchor="w", padx=22, pady=(0, 14))

        corps = tk.Frame(dlg, bg="#0d1117")
        corps.pack(fill="both", expand=True, padx=22)

        # État par document : {clé: {"var": BooleanVar, "fichier": str|None,
        #                            "lbl_fichier": Label, "autre_var": StringVar}}
        etat = {}

        for cle, libelle in DOCS:
            var = tk.BooleanVar(value=False)
            autre_var = tk.StringVar(value="")
            etat[cle] = {"var": var, "fichier": None,
                         "lbl_fichier": None, "autre_var": autre_var}

            # Conteneur pour la ligne complète du document.
            frame_doc = tk.Frame(corps, bg="#0d1117")
            frame_doc.pack(fill="x", pady=(0, 10))

            # Case à cocher + libellé.
            cb = tk.Checkbutton(
                frame_doc, text=libelle, variable=var,
                bg="#0d1117", fg="#e6edf3", selectcolor="#0d1117",
                activebackground="#0d1117", activeforeground="#e6edf3",
                font=("Segoe UI", 12, "bold"), anchor="w",
            )
            cb.pack(anchor="w")

            # Champ texte "Autre document".
            if cle == "autre":
                frame_autre = tk.Frame(frame_doc, bg="#0d1117")
                frame_autre.pack(anchor="w", padx=26, pady=(2, 0))
                tk.Label(frame_autre, text="Précisez le type :",
                         bg="#0d1117", fg="#8b949e",
                         font=("Segoe UI", 10)).pack(side="left")
                tk.Entry(frame_autre, textvariable=autre_var, width=28,
                         font=("Segoe UI", 10)).pack(side="left", padx=(6, 0))

            # Sous-frame bouton + nom fichier (visible quand coché).
            frame_fichier = tk.Frame(frame_doc, bg="#0d1117")
            frame_fichier.pack(anchor="w", padx=26, pady=(4, 0))

            lbl_fich = tk.Label(frame_fichier, text="Aucun fichier joint",
                                bg="#0d1117", fg="#8b949e",
                                font=("Segoe UI", 9), anchor="w")
            etat[cle]["lbl_fichier"] = lbl_fich

            def _joindre(c=cle, lf=lbl_fich):
                p = filedialog.askopenfilename(
                    parent=dlg,
                    title="Choisir un fichier",
                    filetypes=[
                        ("Images et PDF", "*.jpg *.jpeg *.png *.pdf"),
                        ("Images", "*.jpg *.jpeg *.png"),
                        ("PDF", "*.pdf"),
                        ("Tous les fichiers", "*.*"),
                    ],
                )
                if p:
                    etat[c]["fichier"] = p
                    nom_court = os.path.basename(p)
                    if len(nom_court) > 40:
                        nom_court = nom_court[:37] + "..."
                    lf.config(text=nom_court, fg="#7ee787")

            tk.Button(
                frame_fichier,
                text="\U0001f4ce Joindre un fichier (image ou PDF)",
                command=_joindre,
                font=("Segoe UI", 10), pady=4, padx=8,
            ).pack(side="left")
            lbl_fich.pack(side="left", padx=(10, 0))

        # --- Boutons bas ---
        def continuer():
            docs = []
            for cle, libelle in DOCS:
                if etat[cle]["var"].get():
                    if cle == "autre":
                        lab = etat[cle]["autre_var"].get().strip() or "Autre document"
                    else:
                        lab = libelle
                    docs.append({"label": lab, "fichier": etat[cle]["fichier"]})
            resultat["annexes"] = docs
            resultat["valide"] = True
            dlg.destroy()

        def passer():
            resultat["annexes"] = []
            resultat["valide"] = True
            dlg.destroy()

        barre = tk.Frame(dlg, bg="#161b22")
        barre.pack(side="bottom", fill="x")
        tk.Button(barre, text="Continuer", command=continuer,
                  font=("Segoe UI", 11, "bold"), width=16,
                  ).pack(side="right", padx=14, pady=10)
        tk.Button(barre, text="Passer — aucun document", command=passer,
                  font=("Segoe UI", 11), width=24,
                  ).pack(side="right", padx=(0, 6), pady=10)

        dlg.protocol("WM_DELETE_WINDOW", passer)   # croix = Passer
        self.root.wait_window(dlg)
        return resultat["annexes"]

    def _enregistrer(self, infos, resume=None, annexes=None):
        """Produit un fichier Word .docx (avec fallback .txt si python-docx echoue).
        Structure : titre + tableau patient + resume edite + transcription integrale
        + documents annexes. Renvoie True si ecrit, False si la boite est annulee."""
        now = datetime.datetime.now()
        date_str  = now.strftime("%d/%m/%Y")
        heure_str = now.strftime("%Hh%M")
        nom    = infos["nom"]
        prenom = infos["prenom"]

        nom_defaut = "%s_%s_%s_%s.docx" % (
            nettoyer_nom_fichier(nom.upper()),
            nettoyer_nom_fichier(prenom),
            now.strftime("%Y-%m-%d"),
            now.strftime("%Hh%M"),
        )
        initialdir = self.dossier_sauvegarde \
            if (self.dossier_sauvegarde and os.path.isdir(self.dossier_sauvegarde)) else None

        chemin = filedialog.asksaveasfilename(
            parent=self.root,
            title="Enregistrer le compte-rendu",
            defaultextension=".docx",
            initialfile=nom_defaut,
            initialdir=initialdir,
            filetypes=[("Document Word", "*.docx"), ("Tous les fichiers", "*.*")],
        )
        if not chemin:
            return False

        # Tente l'enregistrement docx ; repli txt si python-docx indisponible.
        try:
            _ecrire_docx(chemin, infos, now, resume, self.entries, annexes=annexes)
        except Exception:
            # Fallback .txt (ne jamais perdre la transcription).
            chemin_txt = re.sub(r"\.docx$", ".txt", chemin, flags=re.IGNORECASE)
            if not chemin_txt.lower().endswith(".txt"):
                chemin_txt += ".txt"
            try:
                _ecrire_txt_secours(chemin_txt, infos, now, resume, self.entries)
                message_simple(
                    "Enregistrement Word impossible",
                    "Le document Word n'a pas pu être créé.\n\n"
                    "La transcription a été sauvegardée en texte brut :\n"
                    + chemin_txt)
            except Exception:
                message_simple(
                    "Enregistrement impossible",
                    "Le fichier n'a pas pu être enregistré à cet endroit.\n\n"
                    "Choisissez un autre dossier (par exemple le Bureau).",
                    genre="error")
                return False

        try:
            dossier = os.path.dirname(chemin)
            self.dossier_sauvegarde = dossier
            cfg = charger_config()
            cfg["dossier_sauvegarde"] = dossier
            sauver_config(cfg)
        except Exception:
            pass
        return True


# ----------------------------- MAIN (tkinter legacy) -------------------------

def _main_tk():
    """Point d'entree tkinter (conserve pour --tk et tests en dev)."""
    choix = choisir_peripheriques()
    if choix is None:
        return
    loopback, micro, dossier = choix

    banniere = None
    if micro is None:
        banniere = ("Aucun micro détecté — seule la voix du patient sera transcrite. "
                    "Branchez un micro puis relancez pour transcrire aussi le médecin.")
        display_queue.put(("AVIS", banniere))
    else:
        display_queue.put(("INFO", "Micro (médecin) : " + micro.name))
    display_queue.put(("INFO", "Sortie (patient) : " + loopback.name))

    threads = [
        threading.Thread(target=transcrire, daemon=True),
        threading.Thread(target=capturer, args=(lambda: loopback, "Patient"), daemon=True),
    ]
    if micro is not None:
        threads.append(
            threading.Thread(target=capturer, args=(lambda: micro, "Medecin"), daemon=True))
    for t in threads:
        t.start()

    root = tk.Tk()
    Overlay(root, dossier_sauvegarde=dossier, banniere=banniere)
    try:
        root.mainloop()
    finally:
        stop_event.set()
        time.sleep(0.3)


# ----------------------------- GESTIONNAIRE DE TÉLÉCHARGEMENT ----------------

# État partagé entre le thread de téléchargement et l'Api (JSON-sérialisable).
# ---- Logger fichier pour diagnostic des téléchargements ---------------
import logging as _logging
import os as _os
_log_path = _os.path.join(_os.environ.get('APPDATA', _os.path.expanduser('~')), 'Echo', 'download.log')
_os.makedirs(_os.path.dirname(_log_path), exist_ok=True)
_logging.basicConfig(
    filename=_log_path, level=_logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    force=True)

_download_state = {
    "whisper": {"downloaded": 0, "total": 1_700_000_000, "speed": 0.0,
                "done": False, "error": None},
    "qwen":    {"downloaded": 0, "total": 2_050_000_000, "speed": 0.0,
                "done": False, "error": None},
    "running": False,
}
_download_lock = threading.Lock()


def _taille_dossier(chemin):
    """Taille totale en octets de tous les fichiers d'un dossier."""
    total = 0
    try:
        for root, _, files in os.walk(chemin):
            for f in files:
                try:
                    total += os.path.getsize(os.path.join(root, f))
                except OSError:
                    pass
    except OSError:
        pass
    return total


def _taille_fichier(chemin):
    try:
        return os.path.getsize(chemin) if os.path.isfile(chemin) else 0
    except OSError:
        return 0


def _progres_whisper_loop(dest_dir, stop_dl):
    """Thread de polling pour la progression Whisper."""
    prev_dl   = 0
    prev_time = time.time()
    while not stop_dl.is_set():
        dl = _taille_dossier(dest_dir)
        now = time.time()
        dt  = now - prev_time
        if dt > 0:
            speed = (dl - prev_dl) / dt  # octets/s
        else:
            speed = 0.0
        prev_dl   = dl
        prev_time = now
        with _download_lock:
            s = _download_state["whisper"]
            if not s["done"] and not s["error"]:
                s["downloaded"] = dl
                s["speed"]      = round(speed / 1_000_000, 1)  # Mo/s
        time.sleep(0.5)


def _progres_qwen_loop(dest_file, stop_dl):
    """Thread de polling pour la progression Qwen."""
    prev_dl   = 0
    prev_time = time.time()
    while not stop_dl.is_set():
        dl  = _taille_fichier(dest_file)
        now = time.time()
        dt  = now - prev_time
        if dt > 0:
            speed = (dl - prev_dl) / dt
        else:
            speed = 0.0
        prev_dl   = dl
        prev_time = now
        with _download_lock:
            s = _download_state["qwen"]
            if not s["done"] and not s["error"]:
                s["downloaded"] = dl
                s["speed"]      = round(speed / 1_000_000, 1)
        time.sleep(0.5)


def _dl_whisper():
    """Télécharge le modèle Whisper depuis HuggingFace (avec reprise)."""
    import logging
    try:
        logging.debug("_dl_whisper: début")
        from huggingface_hub import snapshot_download
        dest = os.path.join(models_dir(), MODELE_WHISPER_DIR)
        logging.debug("_dl_whisper: dest = %s", dest)
        os.makedirs(dest, exist_ok=True)
        stop_dl = threading.Event()
        t_prog  = threading.Thread(target=_progres_whisper_loop,
                                   args=(dest, stop_dl), daemon=True)
        t_prog.start()
        logging.debug("_dl_whisper: snapshot_download...")
        snapshot_download(
            repo_id="deepdml/faster-whisper-large-v3-turbo-ct2",
            local_dir=dest,
            resume_download=True,
            ignore_patterns=["*.msgpack", "*.h5", "flax_model*",
                             "tf_model*", "*.ot", "README*"],
        )
        logging.debug("_dl_whisper: terminé avec succès")
        with _download_lock:
            _download_state["whisper"]["done"]       = True
            _download_state["whisper"]["downloaded"] = _download_state["whisper"]["total"]
            _download_state["whisper"]["speed"]      = 0.0
    except Exception as exc:
        logging.exception("_dl_whisper: exception")
        with _download_lock:
            _download_state["whisper"]["error"] = (
                "Téléchargement Whisper interrompu. "
                "Vérifiez votre connexion internet et réessayez. "
                "(" + str(exc)[:80] + ")")
    finally:
        stop_dl.set()


def _dl_qwen():
    """Télécharge le modèle Qwen depuis HuggingFace (avec reprise)."""
    import logging
    try:
        logging.debug("_dl_qwen: début")
        from huggingface_hub import hf_hub_download
        dest_dir  = models_dir()
        dest_file = os.path.join(dest_dir, RESUME_GGUF)
        logging.debug("_dl_qwen: dest_file = %s", dest_file)
        os.makedirs(dest_dir, exist_ok=True)
        stop_dl = threading.Event()
        t_prog  = threading.Thread(target=_progres_qwen_loop,
                                   args=(dest_file, stop_dl), daemon=True)
        t_prog.start()
        logging.debug("_dl_qwen: hf_hub_download...")
        hf_hub_download(
            repo_id="bartowski/Qwen2.5-3B-Instruct-GGUF",
            filename=RESUME_GGUF,
            local_dir=dest_dir,
            resume_download=True,
        )
        logging.debug("_dl_qwen: terminé avec succès")
        with _download_lock:
            _download_state["qwen"]["done"]       = True
            _download_state["qwen"]["downloaded"] = _download_state["qwen"]["total"]
            _download_state["qwen"]["speed"]      = 0.0
    except Exception as exc:
        logging.exception("_dl_qwen: exception")
        with _download_lock:
            _download_state["qwen"]["error"] = (
                "Téléchargement Qwen interrompu. "
                "Vérifiez votre connexion internet et réessayez. "
                "(" + str(exc)[:80] + ")")
    finally:
        stop_dl.set()


def _run_downloads(models_to_dl):
    """Télécharge séquentiellement les modèles demandés."""
    import logging
    logging.debug("_run_downloads démarré, modèles: %s", models_to_dl)
    try:
        with _download_lock:
            _download_state["running"] = True
        logging.debug("running=True")
        if "whisper" in models_to_dl:
            _dl_whisper()
        if "qwen" in models_to_dl:
            _dl_qwen()
        logging.debug("_run_downloads terminé avec succès")
    except Exception:
        logging.exception("_run_downloads: exception")
    finally:
        logging.debug("_run_downloads finally: running=False")
        with _download_lock:
            _download_state["running"] = False


# ----------------------------- API PYWEBVIEW ---------------------------------

class Api:
    """Pont JS ↔ Python exposé via js_api= de pywebview.
    Toutes les méthodes sont appelées depuis un thread worker de pywebview
    (thread-safe requis). Elles doivent être synchrones et retourner des
    types JSON-sérialisables (dict, list, str, int, float, bool, None)."""

    def __init__(self):
        self._window      = None  # fenêtre générique (compatibilité Phase 1)
        self._main_win    = None  # fenêtre principale
        self._overlay_win = None  # overlay transcript
        self._lock        = threading.Lock()
        self._entries     = []    # (timestamp, label, texte)
        self._infos       = None
        self._resume      = None
        self._annexes     = []
        self._dossier_sauvegarde = None
        self._started     = False
        self._save_done   = False
        self._start_time  = None
        self._resume_status = "idle"
        self._resume_text   = None
        self._webview_mod   = None  # module webview, disponible après start

    # ---- Polling -------------------------------------------------------

    def get_updates(self):
        """Vide la display_queue et retourne les nouvelles entrées.
        Appelé toutes les ~200 ms depuis le JS.
        Retourne [{type, texte, timestamp}, ...]."""
        items = []
        while True:
            try:
                label, texte = display_queue.get_nowait()
            except queue.Empty:
                break
            ts = datetime.datetime.now().strftime("%H:%M:%S")
            if label in ("Medecin", "Patient"):
                with self._lock:
                    self._entries.append((ts, label, texte))
            items.append({"type": label, "texte": texte, "timestamp": ts})
        return items

    def get_speaking_status(self):
        """État VAD temps réel pour l'indicateur « qui parle ».
        Appelé toutes les ~120 ms depuis le JS."""
        return {
            "medecin": _speaking_now.get("medecin", False),
            "patient": _speaking_now.get("patient", False),
            "active":  not stop_event.is_set(),
        }

    # ---- Périphériques -------------------------------------------------

    def get_devices(self):
        """Retourne les périphériques disponibles avec pré-sélections.
        Retourne {outputs, mics, selected_output, selected_mic}."""
        cfg     = charger_config()
        outputs = [str(s.name) for s in lister_sorties()]
        mics    = [str(m.name) for m in lister_micros()]

        # Pré-sélection sortie : config > PREFERRED_OUTPUT > défaut système.
        presel_out = cfg.get("sortie") if cfg.get("sortie") in outputs else None
        if not presel_out and PREFERRED_OUTPUT:
            for n in outputs:
                if PREFERRED_OUTPUT.lower() in n.lower():
                    presel_out = n
                    break
        if not presel_out:
            d = nom_sortie_defaut()
            presel_out = d if d in outputs else (outputs[0] if outputs else "")

        # Pré-sélection micro.
        presel_mic = cfg.get("micro") if cfg.get("micro") in mics else None
        if not presel_mic:
            d = nom_micro_defaut()
            presel_mic = d if d in mics else (mics[0] if mics else "")

        self._dossier_sauvegarde = cfg.get("dossier_sauvegarde")
        return {
            "outputs":         outputs,
            "mics":            mics,
            "selected_output": presel_out,
            "selected_mic":    presel_mic,
        }

    # ---- Démarrage -----------------------------------------------------

    def start(self, mic_name, output_name):
        """Résout les périphériques, sauvegarde la config et démarre les threads.
        Retourne {ok, loopback?, mic?, warning?} ou {ok:false, error}."""
        if self._started:
            return {"ok": False, "error": "Déjà démarré."}

        loopback = resoudre_loopback(output_name) or loopback_defaut()
        if loopback is None:
            return {"ok": False,
                    "error": ("Aucune sortie audio disponible. "
                              "Branchez un casque ou des haut-parleurs, "
                              "puis relancez Écho.")}

        micro = resoudre_micro(mic_name) if mic_name else None

        # Persistance config.
        cfg = charger_config()
        cfg["sortie"] = output_name
        cfg["micro"]  = mic_name or ""
        sauver_config(cfg)

        threads = [
            threading.Thread(target=transcrire, daemon=True),
            threading.Thread(target=capturer,
                             args=(lambda: loopback, "Patient"), daemon=True),
        ]
        if micro is not None:
            threads.append(
                threading.Thread(target=capturer,
                                 args=(lambda: micro, "Medecin"), daemon=True))
        for t in threads:
            t.start()

        self._started = True
        result = {"ok": True, "loopback": loopback.name}
        if micro is None:
            result["warning"] = (
                "Aucun micro détecté — seule la voix du patient sera transcrite. "
                "Branchez un micro puis relancez pour transcrire aussi le médecin.")
        else:
            result["mic"] = micro.name
        return result

    # ---- Fermeture (Phase 1 : simple ; Phase 2 : flux de sauvegarde) ---

    def quit_app(self):
        """Phase 1 : arrête les threads et ferme la fenêtre."""
        stop_event.set()
        if self._window:
            self._window.destroy()
        return {"ok": True}

    # ---- Phase 2+ : stubs exposés dès maintenant -----------------------

    def save_patient(self, nom, prenom, ddn, motif):
        """Enregistre les infos patient (Phase 2 : déclenche le flux résumé)."""
        self._infos = {"nom": nom, "prenom": prenom, "naissance": ddn, "motif": motif}
        return {"ok": True}

    def get_resume(self):
        """Retourne le résumé généré (Phase 2). Stub Phase 1."""
        return {"texte": self._resume or ""}

    def validate_resume(self, texte_corrige):
        """Valide le résumé édité et déclenche la sauvegarde du .docx (Phase 2)."""
        self._resume = texte_corrige
        return {"ok": True}

    def set_annexes(self, annexes):
        """Enregistre la liste des documents annexes (Phase 2)."""
        self._annexes = annexes or []
        return {"ok": True}

    def pick_file(self):
        """Ouvre un sélecteur de fichier natif via pywebview (Phase 2)."""
        if not self._window:
            return {"path": None}
        try:
            import webview as _wv
            paths = self._window.create_file_dialog(
                _wv.OPEN_DIALOG,
                allow_multiple=False,
                file_types=("Images et PDF (*.jpg;*.jpeg;*.png;*.pdf)",
                            "Tous les fichiers (*.*)"),
            )
            return {"path": paths[0] if paths else None}
        except Exception:
            return {"path": None}

    def cancel_save(self):
        """Annule le flux de sauvegarde côté overlay (JS annule le modal)."""
        return {"ok": True}

    def navigate(self, target):
        """Charge une nouvelle page dans une fenêtre (usage interne)."""
        win = self._overlay_win or self._main_win or self._window
        if not win:
            return {"ok": False}
        ui_dir = ressource("ui")
        if not target.endswith(".html"):
            target += ".html"
        win.load_url(os.path.join(ui_dir, target))
        return {"ok": True}

    # ===== ÉTAT GLOBAL / ONBOARDING ==========================================

    def get_app_state(self):
        """Retourne l'état de l'appli : onboarding fait, nom médecin, etc."""
        cfg = charger_config()
        return {
            "onboarding_done":   bool(cfg.get("doctor_name")),
            "doctor_name":       cfg.get("doctor_name", ""),
            "save_folder":       cfg.get("dossier_sauvegarde", ""),
            "gain_patient":      cfg.get("gain_patient", 1.0),
            "gain_mic":          cfg.get("gain_mic", 1.0),
        }

    def complete_onboarding(self, doctor_name, mic_name, output_name):
        """Valide l'onboarding et sauvegarde la config complète."""
        if not doctor_name.strip():
            return {"ok": False, "error": "Le nom du médecin est obligatoire."}
        cfg = charger_config()
        cfg["doctor_name"] = doctor_name.strip()
        cfg["micro"]  = mic_name
        cfg["sortie"] = output_name
        sauver_config(cfg)
        return {"ok": True}

    # ===== SETTINGS ==========================================================

    def get_settings(self):
        """Config + périphériques (pour la page Paramètres)."""
        cfg  = charger_config()
        devs = self.get_devices()
        return {
            "doctor_name":   cfg.get("doctor_name", ""),
            "save_folder":   cfg.get("dossier_sauvegarde", ""),
            "gain_patient":  cfg.get("gain_patient", 1.0),
            "gain_mic":      cfg.get("gain_mic", 1.0),
            **devs,
        }

    def save_settings(self, data):
        """Sauvegarde les paramètres modifiés."""
        cfg = charger_config()
        for k in ("doctor_name", "save_folder"):
            if k in data and data[k] is not None:
                key = "dossier_sauvegarde" if k == "save_folder" else k
                cfg[key] = data[k]
        if "gain_patient" in data:
            v = max(0.0, min(1.5, float(data["gain_patient"])))
            cfg["gain_patient"] = v
            self.set_volume_patient(v)
        if "gain_mic" in data:
            v = max(0.0, min(1.5, float(data["gain_mic"])))
            cfg["gain_mic"] = v
            self.set_volume_mic(v)
        if "sortie" in data: cfg["sortie"] = data["sortie"]
        if "micro"  in data: cfg["micro"]  = data["micro"]
        sauver_config(cfg)
        return {"ok": True}

    def pick_folder(self):
        """Sélecteur de dossier natif (depuis n'importe quelle fenêtre)."""
        win = self._main_win or self._overlay_win or self._window
        if not win:
            return {"path": None}
        try:
            import webview as _wv
            paths = win.create_file_dialog(_wv.FOLDER_DIALOG)
            return {"path": paths[0] if paths else None}
        except Exception:
            return {"path": None}

    # ===== VOLUME ============================================================

    def set_volume_patient(self, ratio):
        global _GAIN_PATIENT
        with _gain_lock:
            _GAIN_PATIENT = max(0.0, min(1.5, float(ratio)))
        return {"ok": True}

    def set_volume_mic(self, ratio):
        global _GAIN_MIC
        with _gain_lock:
            _GAIN_MIC = max(0.0, min(1.5, float(ratio)))
        return {"ok": True}

    # ===== DÉMARRAGE / FIN DE CONSULTATION ===================================

    def begin_consultation(self, mic_name, output_name):
        """Lance la consultation depuis la fenêtre principale."""
        # Réinitialise l'état pour une nouvelle consultation.
        stop_event.clear()
        with self._lock:
            self._entries.clear()
        self._save_done  = False
        self._infos      = None
        self._resume     = None
        self._annexes    = []
        self._resume_status = "idle"
        self._resume_text   = None
        self._start_time    = datetime.datetime.now()

        result = self.start(mic_name, output_name)
        if result.get("ok"):
            if self._main_win:
                self._main_win.minimize()
            if self._overlay_win:
                self._overlay_win.show()
                self._overlay_win.restore()
        return result

    def end_consultation(self):
        """Ferme l'overlay et restaure la fenêtre principale."""
        self._save_done = True
        if self._overlay_win:
            self._overlay_win.hide()
        if self._main_win:
            self._main_win.restore()
            self._main_win.evaluate_js("onConsultationEnded()")
        stop_event.set()
        return {"ok": True}

    def request_quit(self):
        """Déclenche le flux de fermeture (depuis le bouton Quitter ou la croix).
        Retourne {needs_save: true} si des entrées non sauvegardées existent,
        {needs_save: false} si on peut fermer directement."""
        with self._lock:
            has_entries = bool(self._entries)
        if has_entries and not getattr(self, "_save_done", False):
            return {"needs_save": True}
        self._close_all()
        return {"needs_save": False}

    def force_quit(self):
        """Fermeture sans sauvegarde (confirmée par l'utilisateur)."""
        self._close_all()
        return {"ok": True}

    def _close_all(self):
        stop_event.set()
        for w in (self._overlay_win, self._main_win, self._window):
            if w:
                try:
                    w.destroy()
                except Exception:
                    pass

    # ===== FLUX DE SAUVEGARDE (orchestré depuis le JS de l'overlay) ==========

    def generate_resume_async(self):
        """Lance la génération du résumé dans un thread.
        JS pollera get_resume_status() pour la progression."""
        self._resume_status = "loading"
        self._resume_text   = None

        def worker():
            try:
                llm = charger_modele_resume()
                self._resume_status = "generating"
                with self._lock:
                    entries = list(self._entries)
                transcript = "\n".join(
                    "[%s] %s : %s" % (h, LOCUTEUR_FICHIER.get(loc, loc), t)
                    for h, loc, t in entries)
                self._resume_text   = generer_resume(llm, transcript)
                self._resume_status = "done"
            except Exception:
                self._resume_status = "error"

        threading.Thread(target=worker, daemon=True).start()
        return {"ok": True}

    def get_resume_status(self):
        return {
            "status": getattr(self, "_resume_status", "idle"),
            "texte":  getattr(self, "_resume_text",   None) or "",
        }

    def open_save_dialog(self, default_filename):
        """Ouvre la boîte 'Enregistrer sous' .docx via pywebview."""
        win = self._overlay_win or self._window
        if not win:
            return {"path": None}
        try:
            import webview as _wv
            cfg     = charger_config()
            initdir = cfg.get("dossier_sauvegarde") or ""
            paths   = win.create_file_dialog(
                _wv.SAVE_DIALOG,
                save_filename=default_filename,
                directory=initdir if os.path.isdir(initdir) else "",
                file_types=("Document Word (*.docx)", "Tous les fichiers (*.*)")
            )
            return {"path": paths[0] if paths else None}
        except Exception:
            return {"path": None}

    def perform_save(self, file_path, resume_text, annexes):
        """Écrit le .docx, enregistre dans consultations.json."""
        if not self._infos:
            return {"ok": False, "error": "Informations patient manquantes."}
        now = getattr(self, "_start_time", None) or datetime.datetime.now()
        with self._lock:
            entries = list(self._entries)
        try:
            _ecrire_docx(file_path, self._infos, now,
                         resume_text or None, entries, annexes=annexes or [])
        except Exception as exc:
            # Fallback txt.
            txt_path = re.sub(r"\.docx$", ".txt", file_path, flags=re.IGNORECASE)
            try:
                _ecrire_txt_secours(txt_path, self._infos, now, resume_text, entries)
                file_path = txt_path
            except Exception:
                return {"ok": False, "error": str(exc)}

        # Mémorise le dossier.
        dossier = os.path.dirname(file_path)
        cfg = charger_config()
        cfg["dossier_sauvegarde"] = dossier
        sauver_config(cfg)

        # Ajoute au journal.
        dur = int((datetime.datetime.now() - now).total_seconds() / 60)
        ajouter_consultation({
            "id":           str(uuid.uuid4()),
            "date":         now.isoformat(),
            "patient":      self._infos,
            "summary":      resume_text or "",
            "file_path":    file_path,
            "duration_min": dur,
        })
        return {"ok": True}

    # ===== HISTORIQUE =========================================================

    def get_consultations(self):
        return charger_consultations()

    def delete_consultation(self, cid):
        """Retire l'entrée d'id `cid` de consultations.json.

        Le fichier .docx sur le disque est conservé.
        """
        try:
            supprimer_consultation(cid)
            return {"ok": True}
        except (PermissionError, OSError):
            return {"ok": False,
                    "error": "Fichier temporairement inaccessible. Réessayez."}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    def delete_consultation_with_file(self, cid):
        """Retire l'entrée d'id `cid` et supprime aussi le .docx s'il existe.

        Si le fichier n'existe pas/plus, l'entrée est quand même retirée
        sans erreur.
        """
        try:
            record = supprimer_consultation(cid)
        except (PermissionError, OSError):
            return {"ok": False,
                    "error": "Fichier temporairement inaccessible. Réessayez."}
        except Exception as e:
            return {"ok": False, "error": str(e)}
        try:
            fp = (record or {}).get("file_path")
            if fp and os.path.isfile(fp):
                os.remove(fp)
        except Exception:
            # Jamais d'erreur visible si le fichier a déjà été
            # déplacé/supprimé manuellement ou est verrouillé.
            pass
        return {"ok": True}

    # ===== FICHIERS (OS) ======================================================

    def open_file_os(self, path):
        try:
            os.startfile(path)
            return {"ok": True}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    def open_folder_os(self, path):
        try:
            subprocess.Popen(["explorer", "/select,", path])
            return {"ok": True}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    def file_exists(self, path):
        return {"exists": bool(path and os.path.isfile(path))}

    def get_patient_infos(self):
        """Retourne les infos patient en cours (pour le nom de fichier par défaut)."""
        return self._infos or {}

    # ===== MODÈLES — VÉRIFICATION ET TÉLÉCHARGEMENT ==========================

    def check_models(self):
        """Vérifie si les modèles sont présents et valides."""
        return {"whisper_ok": whisper_ok(), "qwen_ok": qwen_ok()}

    def start_downloads(self):
        """Lance le(s) téléchargement(s) manquant(s) dans un thread daemon."""
        import logging
        try:
            logging.debug("start_downloads() appelé")
            with _download_lock:
                if _download_state["running"]:
                    logging.debug("déjà en cours")
                    return {"ok": True, "msg": "already running"}
                if not whisper_ok():
                    _download_state["whisper"]["downloaded"] = 0
                    _download_state["whisper"]["done"]       = False
                    _download_state["whisper"]["error"]      = None
                if not qwen_ok():
                    _download_state["qwen"]["downloaded"] = 0
                    _download_state["qwen"]["done"]       = False
                    _download_state["qwen"]["error"]      = None

            to_dl = []
            if not whisper_ok(): to_dl.append("whisper")
            if not qwen_ok():    to_dl.append("qwen")

            with _download_lock:
                if "whisper" not in to_dl:
                    _download_state["whisper"]["done"] = True
                if "qwen" not in to_dl:
                    _download_state["qwen"]["done"] = True

            logging.debug("modèles à télécharger: %s", to_dl)
            if to_dl:
                threading.Thread(target=_run_downloads, args=(to_dl,), daemon=True).start()
            logging.debug("start_downloads() retourne ok")
            return {"ok": True}
        except Exception as e:
            logging.exception("start_downloads: exception")
            return {"ok": False, "error": str(e)}

    def get_download_progress(self):
        """Retourne l'état complet des téléchargements (JSON-serializable)."""
        with _download_lock:
            import copy
            return copy.deepcopy(_download_state)

    def retry_download(self, model):
        """Remet error=None et relance le téléchargement d'un modèle spécifique."""
        if model not in ("whisper", "qwen"):
            return {"ok": False, "error": "modèle inconnu"}
        with _download_lock:
            if _download_state["running"]:
                return {"ok": False, "error": "Téléchargement déjà en cours."}
            _download_state[model]["error"]      = None
            _download_state[model]["done"]       = False
            _download_state[model]["downloaded"] = 0
        threading.Thread(target=_run_downloads, args=([model],), daemon=True).start()
        return {"ok": True}


# ----------------------------- MAIN (pywebview) ------------------------------

def _main_webview():
    """Point d'entrée pywebview — deux fenêtres distinctes."""
    import webview

    api    = Api()
    ui_dir = ressource("ui")

    # Fenêtre 1 : application principale (accueil / paramètres / historique).
    main_win = webview.create_window(
        "Écho",
        os.path.join(ui_dir, "main_window.html"),
        js_api=api,
        width=870, height=660,
        min_size=(720, 520),
        background_color="#F7F8FA",
    )
    # Fenêtre 2 : overlay de transcription (toujours au premier plan).
    overlay_win = webview.create_window(
        "Écho — Consultation en cours",
        os.path.join(ui_dir, "index.html"),
        js_api=api,
        width=560, height=720,
        min_size=(480, 520),
        on_top=True,
        background_color="#080E1C",
    )

    api._main_win    = main_win
    api._overlay_win = overlay_win
    api._window      = overlay_win   # compatibilité Phase 1

    # Intercepter la fermeture de l'overlay via la croix Windows.
    def on_overlay_closing():
        with api._lock:
            has_entries = bool(api._entries)
        if has_entries and not api._save_done:
            overlay_win.evaluate_js("handleWindowClose()")
            return False
        return True

    overlay_win.events.closing += on_overlay_closing

    def on_start():
        api._webview_mod = webview
        overlay_win.hide()   # L'overlay est caché jusqu'au début d'une consultation.

    webview.start(func=on_start, debug=("--dev" in sys.argv))
    stop_event.set()
    time.sleep(0.2)


def main():
    if "--tk" in sys.argv:
        _main_tk()
    else:
        _main_webview()


def selftest():
    """Validation headless de l'environnement (installeur léger).
    Vérifie les imports critiques et l'état des modèles sans les télécharger.
    Écrit le résultat dans un log + code de sortie."""
    import traceback
    log = os.path.join(tempfile.gettempdir(), "echo_selftest.log")
    try:
        import av          # noqa: F401
        import ctranslate2 # noqa: F401
        import onnxruntime # noqa: F401
        from huggingface_hub import hf_hub_download  # noqa: F401 — requis pour dl
        frozen = getattr(sys, "frozen", False)
        w_ok   = whisper_ok()
        q_ok   = qwen_ok()
        models_info = "whisper=%s qwen=%s" % (w_ok, q_ok)
        # Charge le modèle Whisper si disponible, sinon valide juste les imports.
        if w_ok:
            source = chemin_modele()
            model  = WhisperModel(source, device=DEVICE, compute_type=COMPUTE_TYPE)
            segs, _ = model.transcribe(np.zeros(SAMPLE_RATE, dtype=np.float32),
                                       language=LANGUAGE, vad_filter=True, beam_size=1)
            list(segs)
            models_info += " whisper_loaded=True"
        msg  = ("SELFTEST OK | frozen=%s | meipass=%s | models=%s | models_dir=%s"
                % (frozen, getattr(sys, "_MEIPASS", "-"),
                   models_info, models_dir()))
        code = 0
    except Exception:
        msg  = "SELFTEST FAIL\n" + traceback.format_exc()
        code = 1
    try:
        with open(log, "w", encoding="utf-8") as f:
            f.write(msg + "\n")
    except Exception:
        pass
    print(msg)
    sys.exit(code)


if __name__ == "__main__":
    if "--selftest" in sys.argv:
        selftest()
    else:
        main()
