"""Persistance d'Écho : historique JSON + génération du document Word.

Module pur, sans dépendance à l'état global de l'application : toutes les
fonctions reçoivent en paramètre le chemin du fichier et les données. Cela
permet de les tester isolément (cf. tests/test_storage.py).

Contenu :
  - Historique consultations.json (lecture/écriture, suppression, mise à jour
    du résumé) avec retry sur verrou Windows (OneDrive/antivirus/indexeur).
  - Génération du .docx (python-docx) et fallback texte.
"""

import json
import os
import re
import time

# Libellé accentué écrit dans le fichier (l'interne reste sans accent).
LOCUTEUR_FICHIER = {"Medecin": "Médecin", "Patient": "Patient"}

# Couleurs pour la section transcription.
_BLEU_PATIENT = "2E74B5"   # bleu sobre (hex sans #)
_GRIS_TITRE   = "404040"   # gris foncé pour Heading TRANSCRIPTION


# ============================ HISTORIQUE (JSON) ==============================

def charger_consultations(chemin):
    """Historique des consultations depuis `chemin`. Renvoie [] si absent/corrompu."""
    try:
        with open(chemin, encoding="utf-8") as f:
            d = json.load(f)
        return d if isinstance(d, list) else []
    except Exception:
        return []


def ajouter_consultation(chemin, record):
    """Ajoute un enregistrement en tête de l'historique (best-effort)."""
    try:
        consultations = charger_consultations(chemin)
        consultations.insert(0, record)
        os.makedirs(os.path.dirname(chemin), exist_ok=True)
        with open(chemin, "w", encoding="utf-8") as f:
            json.dump(consultations, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def _lire_liste(chemin):
    """Lit la liste JSON ; [] si absente. Lève sur verrou (géré par l'appelant)."""
    try:
        with open(chemin, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except FileNotFoundError:
        return []


def supprimer_consultation(chemin, cid):
    """Retire l'entrée d'id `cid` de l'historique et réécrit le fichier.

    Lecture et écriture se font toujours via `with open(...)` : aucun handle
    n'est conservé entre les appels. En cas de verrou Windows transitoire
    (OneDrive, antivirus, indexeur), on réessaie jusqu'à 3 fois.

    Renvoie l'enregistrement supprimé (dict) si trouvé, sinon None.
    Lève PermissionError/OSError si le fichier reste inaccessible après les
    tentatives.
    """
    os.makedirs(os.path.dirname(chemin), exist_ok=True)
    derniere_exc = None
    for attempt in range(3):
        try:
            data = _lire_liste(chemin)
            supprime = None
            restantes = []
            for c in data:
                if supprime is None and isinstance(c, dict) and c.get("id") == cid:
                    supprime = c
                else:
                    restantes.append(c)
            with open(chemin, "w", encoding="utf-8") as f:
                json.dump(restantes, f, ensure_ascii=False, indent=2)
            return supprime
        except (PermissionError, OSError) as exc:
            derniere_exc = exc
            if attempt < 2:
                time.sleep(0.15)
    raise derniere_exc


def supprimer_consultation_avec_fichier(chemin, cid):
    """Retire l'entrée d'id `cid` ET supprime le .docx associé s'il existe.

    Si le fichier est absent/déjà supprimé, l'entrée est quand même retirée
    sans erreur. Lève PermissionError/OSError si le JSON reste inaccessible.
    Renvoie l'enregistrement supprimé (dict) ou None.
    """
    record = supprimer_consultation(chemin, cid)
    try:
        fp = (record or {}).get("file_path")
        if fp and os.path.isfile(fp):
            os.remove(fp)
    except Exception:
        # Jamais d'erreur visible si le fichier a déjà été déplacé/supprimé
        # manuellement ou est verrouillé.
        pass
    return record


def maj_consultation_resume(chemin, cid, resume):
    """Met à jour le champ `summary` de l'entrée d'id `cid` (retry sur verrou).
    Les autres champs de l'entrée restent intacts."""
    os.makedirs(os.path.dirname(chemin), exist_ok=True)
    derniere_exc = None
    for attempt in range(3):
        try:
            data = _lire_liste(chemin)
            for c in data:
                if isinstance(c, dict) and c.get("id") == cid:
                    c["summary"] = resume or ""
                    break
            with open(chemin, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            return
        except (PermissionError, OSError) as exc:
            derniere_exc = exc
            if attempt < 2:
                time.sleep(0.15)
    raise derniere_exc


# ============================ DOCUMENT WORD =================================

def _pt(points):
    """Convertit des points en unités Emu (utilisées par python-docx)."""
    from docx.shared import Pt
    return Pt(points)


def _cm(centimetres):
    from docx.shared import Cm
    return Cm(centimetres)


def _rgb(hex6):
    """Retourne un objet RGBColor depuis une chaine hex sans #."""
    from docx.shared import RGBColor
    r, g, b = int(hex6[0:2], 16), int(hex6[2:4], 16), int(hex6[4:6], 16)
    return RGBColor(r, g, b)


def _forcer_style_heading(style, taille, hex_couleur, gras=True, police="Arial"):
    """Surcharge un style Heading : police Arial, couleur imposee (non bleue)."""
    from docx.shared import Pt, RGBColor
    style.font.name = police
    style.font.size = Pt(taille)
    style.font.bold = gras
    style.font.color.rgb = _rgb(hex_couleur)


def _set_interligne(para, multiple):
    """Fixe l'interligne d'un paragraphe (multiple, ex 1.2)."""
    from docx.enum.text import WD_LINE_SPACING
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = multiple


def _ajouter_run(para, texte, gras=False, italique=False,
                 couleur=None, taille=12, police="Arial"):
    """Ajoute un run formaté dans un paragraphe."""
    from docx.shared import Pt
    r = para.add_run(texte)
    r.font.name = police
    r.font.size = Pt(taille)
    r.bold = gras
    r.italic = italique
    if couleur:
        r.font.color.rgb = _rgb(couleur)
    return r


def ecrire_docx(chemin, infos, now, resume, entries, annexes=None):
    """Construit et ecrit le .docx selon la structure demandee.

    `infos` : dict patient (nom obligatoire ; prenom/naissance/motif optionnels).
    `entries` : liste de (horodatage, locuteur, texte).
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # --- Marges A4 2,5 cm ---
    section = doc.sections[0]
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = section.right_margin = Cm(2.5)
        section.top_margin  = section.bottom_margin = Cm(2.5)
    section = doc.sections[0]   # reference gardee pour la largeur images

    # --- Surcharger Heading 1 (noir) et Heading 2 (noir) ---
    _forcer_style_heading(doc.styles["Heading 1"], 14, "000000")
    _forcer_style_heading(doc.styles["Heading 2"], 13, "000000")

    # 1. Titre principal centré.
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_interligne(titre, 1.2)
    run_titre = titre.add_run("Compte-rendu de consultation")
    run_titre.font.name = "Arial"
    run_titre.font.size = Pt(18)
    run_titre.bold = True
    run_titre.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()   # espace

    # 2. Tableau patient 2 colonnes.
    date_str  = now.strftime("%d/%m/%Y")
    heure_str = now.strftime("%Hh%M")
    # Seul le Nom est obligatoire ; Prénom et Né(e) le sont omis si vides.
    lignes_patient = [("Nom", infos.get("nom") or "—")]
    if (infos.get("prenom") or "").strip():
        lignes_patient.append(("Prénom", infos["prenom"]))
    if (infos.get("naissance") or "").strip():
        lignes_patient.append(("Né(e) le", infos["naissance"]))
    lignes_patient.append(("Date", "%s à %s" % (date_str, heure_str)))
    lignes_patient.append(("Motif", infos.get("motif") or "—"))
    tbl = doc.add_table(rows=len(lignes_patient), cols=2)
    tbl.style = "Table Grid"
    for row_idx, (lbl, val) in enumerate(lignes_patient):
        c0, c1 = tbl.rows[row_idx].cells
        # Largeur colonnes
        c0.width = Cm(4.5)
        c1.width = Cm(12.5)
        # Étiquette en gras
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(lbl)
        r0.bold = True
        r0.font.name = "Arial"
        r0.font.size = Pt(12)
        # Valeur
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.name = "Arial"
        r1.font.size = Pt(12)
        # Couleur de fond légère pour l'étiquette (gris très clair)
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OE
        shd = _OE("w:shd")
        shd.set(_qn("w:val"), "clear")
        shd.set(_qn("w:color"), "auto")
        shd.set(_qn("w:fill"), "F2F2F2")
        c0.paragraphs[0]._p.get_or_add_pPr().append(shd)
    doc.add_paragraph()   # espace après tableau

    # 3. Séparateur horizontal (paragraphe avec bordure bottom).
    sep_para = doc.add_paragraph()
    sep_pPr = sep_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    sep_pPr.append(pBdr)
    doc.add_paragraph()

    # 4. Section RÉSUMÉ (si présent).
    if resume:
        h1_res = doc.add_heading("Résumé de la consultation", level=1)
        _set_interligne(h1_res, 1.2)

        # Bandeau avertissement en italique grisé.
        avert = doc.add_paragraph()
        _set_interligne(avert, 1.2)
        ra = avert.add_run(
            "⚠ Résumé généré automatiquement — à relire et corriger par le médecin.")
        ra.italic = True
        ra.font.name = "Arial"
        ra.font.size = Pt(11)
        ra.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        doc.add_paragraph()

        # Découper le résumé en sections.
        TITRES_RESUME = [
            "Motif :",
            "Observations / points clés :",
            "Traitements et prescriptions évoqués :",
            "Suivi et recommandations :",
        ]
        lignes_res = resume.splitlines()
        # Sauter la première ligne (ENTETE_RESUME déjà dans bandeau).
        debut = 0
        for idx, l in enumerate(lignes_res):
            if l.strip() and not l.startswith("RÉSUMÉ"):
                debut = idx
                break

        titre_courant = None
        puces_courantes = []

        def _flush_section():
            if titre_courant is None:
                return
            h2 = doc.add_heading(titre_courant.rstrip(":"), level=2)
            _set_interligne(h2, 1.2)
            if puces_courantes == ["Non précisé"] or not puces_courantes:
                p = doc.add_paragraph("Non précisé", style="Normal")
                _set_interligne(p, 1.2)
                p.runs[0].font.name = "Arial"
                p.runs[0].font.size = Pt(12)
                p.runs[0].italic = True
            else:
                for puce in puces_courantes:
                    p = doc.add_paragraph(style="List Bullet")
                    _set_interligne(p, 1.2)
                    _ajouter_run(p, puce, taille=12)

        for ligne in lignes_res[debut:]:
            ligne = ligne.strip()
            if not ligne:
                continue
            matched = next((t for t in TITRES_RESUME if ligne.startswith(t)), None)
            if matched:
                _flush_section()
                titre_courant = matched
                puces_courantes = []
            elif ligne == "Non précisé":
                puces_courantes = ["Non précisé"]
            elif ligne.startswith("- "):
                puces_courantes.append(ligne[2:].strip())
            elif ligne.startswith("-"):
                puces_courantes.append(ligne[1:].strip())
        _flush_section()
        doc.add_paragraph()

    # 5. Saut de page.
    doc.add_page_break()

    # 6. Section TRANSCRIPTION INTÉGRALE.
    h1_tr = doc.add_heading("Transcription intégrale", level=1)
    _set_interligne(h1_tr, 1.2)
    h1_tr.runs[0].font.color.rgb = _rgb(_GRIS_TITRE)

    for horodatage, locuteur, texte in entries:
        loc_label = LOCUTEUR_FICHIER.get(locuteur, locuteur)
        p = doc.add_paragraph()
        _set_interligne(p, 1.15)
        est_medecin = (locuteur == "Medecin")
        # Préfixe locuteur + heure.
        prefixe = "%s (%s) : " % (loc_label, horodatage)
        r_pre = p.add_run(prefixe)
        r_pre.font.name = "Arial"
        r_pre.font.size = Pt(11)
        r_pre.bold = est_medecin
        if not est_medecin:
            r_pre.font.color.rgb = _rgb(_BLEU_PATIENT)
        # Texte transcrit.
        r_txt = p.add_run(texte)
        r_txt.font.name = "Arial"
        r_txt.font.size = Pt(11)
        if not est_medecin:
            r_txt.font.color.rgb = _rgb(_BLEU_PATIENT)

    # 7. Section DOCUMENTS ANNEXES (si au moins un document selectionne).
    if annexes:
        _inserer_annexes_docx(doc, annexes, section)

    doc.save(chemin)


def _pdf_vers_image_bytes(chemin_pdf):
    """Convertit la 1re page d'un PDF en PNG (bytes) via PyMuPDF.
    Renvoie les bytes PNG, ou None si PyMuPDF indisponible / erreur."""
    try:
        import fitz
        doc_pdf = fitz.open(chemin_pdf)
        page = doc_pdf[0]
        pix = page.get_pixmap(dpi=150)
        return pix.tobytes("png")
    except Exception:
        return None


def _inserer_annexes_docx(doc, annexes, page_section):
    """Ajoute la section DOCUMENTS ANNEXES apres la transcription.
    annexes : liste de {"label": str, "fichier": str|None}
    Chaque document est independant — un echec n'arrete pas les autres.
    """
    import io
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Séparateur horizontal.
    sep_para = doc.add_paragraph()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    sep_para._p.get_or_add_pPr().append(pBdr)

    h1_ann = doc.add_heading("Documents annexes", level=1)
    _set_interligne(h1_ann, 1.2)
    h1_ann.runs[0].font.color.rgb = _rgb(_GRIS_TITRE)

    # Largeur utile de la page (pour les images).
    larg_page = (page_section.page_width
                 - page_section.left_margin
                 - page_section.right_margin)

    for doc_info in annexes:
        label   = doc_info.get("label", "Document")
        fichier = doc_info.get("fichier")   # peut être None

        h2 = doc.add_heading(label, level=2)
        _set_interligne(h2, 1.2)

        if not fichier:
            # Case cochée mais aucun fichier joint.
            p = doc.add_paragraph()
            _set_interligne(p, 1.2)
            _ajouter_run(p, "Document remis au patient lors de la consultation.",
                         italique=True, couleur="808080", taille=11)
            continue

        ext = os.path.splitext(fichier)[1].lower()

        # --- Image directe (JPG / PNG) ---
        if ext in (".jpg", ".jpeg", ".png"):
            try:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(fichier, width=larg_page)
            except Exception:
                p = doc.add_paragraph()
                _ajouter_run(p, "Image jointe — impossible d'incorporer le fichier.",
                             italique=True, couleur="808080", taille=11)
            continue

        # --- PDF : tentative de conversion via PyMuPDF ---
        if ext == ".pdf":
            img_bytes = _pdf_vers_image_bytes(fichier)
            if img_bytes:
                try:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(
                        io.BytesIO(img_bytes), width=larg_page)
                except Exception:
                    img_bytes = None
            if not img_bytes:
                p = doc.add_paragraph()
                _ajouter_run(
                    p,
                    "Document PDF joint lors de la consultation — "
                    "conserver séparément.",
                    italique=True, couleur="808080", taille=11)
            continue

        # --- Autre format ---
        p = doc.add_paragraph()
        _ajouter_run(p, "Fichier joint : " + os.path.basename(fichier),
                     italique=True, couleur="808080", taille=11)


def ecrire_txt_secours(chemin, infos, now, resume, entries):
    """Fallback .txt si python-docx echoue (la transcription ne doit jamais etre perdue)."""
    sep = "=" * 60
    with open(chemin, "w", encoding="utf-8") as f:
        f.write(sep + "\n")
        entete = ("%s %s" % ((infos.get("nom") or "").upper(),
                             infos.get("prenom") or "")).strip()
        f.write("  CONSULTATION — %s\n" % entete)
        if (infos.get("naissance") or "").strip():
            f.write("  Né(e) le : %s\n" % infos["naissance"])
        f.write("  Date : %s à %s\n" % (now.strftime("%d/%m/%Y"), now.strftime("%Hh%M")))
        f.write("  Motif : %s\n" % (infos.get("motif") or "—"))
        f.write(sep + "\n\n")
        if resume:
            f.write(resume.strip() + "\n\n")
            f.write(sep + "\n  TRANSCRIPTION INTÉGRALE\n" + sep + "\n\n")
        for horodatage, locuteur, texte in entries:
            f.write("[%s] %s : %s\n" % (
                horodatage, LOCUTEUR_FICHIER.get(locuteur, locuteur), texte))
